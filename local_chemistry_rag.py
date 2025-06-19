#!/usr/bin/env python3
"""
Sistema RAG Local para Química
Professora Luciana - ENEM AI Helper

Baseado no sistema de física, adaptado para química
"""

import os
import streamlit as st
from pathlib import Path
from typing import List, Dict, Any, Optional
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain.llms.base import LLM
from groq import Groq
import docx
import PyPDF2
import re

class GroqLLM(LLM):
    """LLM personalizado para DeepSeek R1 Distill via Groq"""
    
    api_key: str
    model_name: str = "deepseek-r1-distill-llama-70b"
    temperature: float = 0.1
    max_tokens: int = 4000
    
    class Config:
        arbitrary_types_allowed = True
    
    def __init__(self, api_key: str, **kwargs):
        super().__init__(api_key=api_key, **kwargs)
    
    @property
    def _llm_type(self) -> str:
        return "groq"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        """Chama a API Groq"""
        try:
            # Adiciona instruções específicas contra mostrar raciocínio
            system_message = """Você é uma professora. NUNCA mostre seu raciocínio interno ou processo de pensamento. Responda DIRETAMENTE como uma professora explicando para uma aluna. NÃO use frases como "Vou analisar", "Preciso calcular", "Vamos pensar", etc."""
            
            # Cria uma nova instância do cliente a cada chamada para evitar cache corrompido
            client = Groq(api_key=self.api_key)
            response = client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                stop=stop
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Erro na API: {str(e)}"

class LocalChemistryRAG:
    """Sistema RAG para Química com vectorstore local"""
    
    def __init__(self, chemistry_folder_path: str = "./quimica"):
        # Configuração adaptativa para cloud/local
        try:
            from cloud_config import get_config
            self.cloud_config = get_config()
            self.chemistry_folder_path = self.cloud_config.get_documents_path("química")
            self.persist_directory = self.cloud_config.get_vectorstore_path("chemistry")
        except ImportError:
            # Fallback para configuração local
            self.cloud_config = None
            self.chemistry_folder_path = chemistry_folder_path
            self.persist_directory = "./chroma_chemistry_db"
            
        self.documents = []
        self.vectorstore = None
        self.retriever = None
        self.rag_chain = None
        self.memory = None
        self.embeddings = None
        
        # Configura embeddings
        self._setup_embeddings()
        
        # Tenta carregar vectorstore existente
        if not self.load_existing_vectorstore():
            # Se não existir, processa documentos se a pasta existir
            if os.path.exists(self.chemistry_folder_path):
                self.process_chemistry_documents()
    
    def _setup_embeddings(self):
        """Configura embeddings multilíngues"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception as e:
            if 'st' in globals():
                st.error(f"Erro ao configurar embeddings: {str(e)}")
            else:
                print(f"Erro ao configurar embeddings: {str(e)}")
            self.embeddings = None
    
    def process_chemistry_documents(self) -> bool:
        """Processa documentos de química e cria vectorstore"""
        try:
            if not os.path.exists(self.chemistry_folder_path):
                if 'st' in globals():
                    st.warning(f"Pasta {self.chemistry_folder_path} não encontrada")
                else:
                    print(f"Pasta {self.chemistry_folder_path} não encontrada")
                return False
            
            chemistry_files = list(Path(self.chemistry_folder_path).glob("*"))
            
            if not chemistry_files:
                if 'st' in globals():
                    st.warning(f"Nenhum arquivo encontrado em {self.chemistry_folder_path}")
                else:
                    print(f"Nenhum arquivo encontrado em {self.chemistry_folder_path}")
                return False
            
            # Processa cada arquivo
            for file_path in chemistry_files:
                if file_path.is_file():
                    docs = self._process_single_file(str(file_path))
                    self.documents.extend(docs)
            
            if not self.documents:
                if 'st' in globals():
                    st.warning("Nenhum documento processado")
                else:
                    print("Nenhum documento processado")
                return False
            
            # Cria vectorstore
            self._create_vectorstore()
            
            if 'st' in globals():
                st.success(f"✅ Processados {len(self.documents)} documentos de química")
            else:
                print(f"✅ Processados {len(self.documents)} documentos de química")
            return True
            
        except Exception as e:
            if 'st' in globals():
                st.error(f"Erro ao processar documentos de química: {str(e)}")
            else:
                print(f"Erro ao processar documentos de química: {str(e)}")
            return False
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Processa um único arquivo"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            # Determina método de extração baseado na extensão
            if file_extension == '.txt':
                content = self._extract_txt_content(file_path)
            elif file_extension == '.docx':
                content = self._extract_docx_content(file_path)
            elif file_extension == '.pdf':
                content = self._extract_pdf_content(file_path)
            else:
                return []
            
            if not content or len(content.strip()) < 50:
                return []
            
            # Cria divisor de texto
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", ".", " ", ""]
            )
            
            # Divide texto em chunks
            chunks = text_splitter.split_text(content)
            
            # Cria documentos
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "filename": Path(file_path).name,
                        "topic": self._extract_topic_from_filename(Path(file_path).name),
                        "chunk_id": i,
                        "subject": "Química"
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            if 'st' in globals():
                st.warning(f"Erro ao processar {file_path}: {str(e)}")
            else:
                print(f"Erro ao processar {file_path}: {str(e)}")
            return []
    
    def _extract_txt_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .txt"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    if isinstance(content, str):
                        content = content.encode('utf-8', errors='replace').decode('utf-8')
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if 'st' in globals():
            st.warning(f"Não foi possível ler arquivo {file_path} com nenhuma codificação")
        else:
            print(f"Não foi possível ler arquivo {file_path} com nenhuma codificação")
        return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .docx"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            return "\n\n".join(content_parts)
        except Exception as e:
            if 'st' in globals():
                st.warning(f"Erro ao processar DOCX {file_path}: {str(e)}")
            else:
                print(f"Erro ao processar DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .pdf"""
        try:
            content_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Verifica se o PDF está criptografado
                if pdf_reader.is_encrypted:
                    # Tenta descriptografar com senha vazia
                    try:
                        pdf_reader.decrypt("")
                    except:
                        if 'st' in globals():
                            st.warning(f"PDF {Path(file_path).name} está protegido por senha - pulando...")
                        else:
                            print(f"PDF {Path(file_path).name} está protegido por senha - pulando...")
                        return ""
                
                # Extrai texto das páginas
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content_parts.append(f"--- Página {page_num + 1} ---\n{text.strip()}")
                    except Exception as e:
                        if 'st' in globals():
                            st.warning(f"Erro na página {page_num + 1} do PDF {Path(file_path).name}: {str(e)}")
                        else:
                            print(f"Erro na página {page_num + 1} do PDF {Path(file_path).name}: {str(e)}")
                        continue
            
            if content_parts:
                return "\n\n".join(content_parts)
            else:
                if 'st' in globals():
                    st.info(f"PDF {Path(file_path).name} processado mas sem texto extraído")
                else:
                    print(f"PDF {Path(file_path).name} processado mas sem texto extraído")
                return ""
                
        except Exception as e:
            error_msg = str(e)
            if "PyCryptodome" in error_msg:
                if 'st' in globals():
                    st.error(f"PDF {Path(file_path).name} requer descriptografia avançada - instale: pip install PyCryptodome")
                else:
                    print(f"PDF {Path(file_path).name} requer descriptografia avançada - instale: pip install PyCryptodome")
            else:
                if 'st' in globals():
                    st.warning(f"Erro ao processar PDF {Path(file_path).name}: {error_msg}")
                else:
                    print(f"Erro ao processar PDF {Path(file_path).name}: {error_msg}")
            return ""
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extrai tópico do nome do arquivo"""
        # Remove extensão
        topic = Path(filename).stem
        
        # Limpa caracteres especiais
        topic = topic.replace("_", " ").replace("-", " ")
        
        # Capitaliza primeira letra
        return topic.title()
    
    def _create_vectorstore(self):
        """Cria ou atualiza o vectorstore"""
        try:
            if not self.embeddings:
                raise Exception("Embeddings não configurados")
            
            # Cria vectorstore ChromaDB
            self.vectorstore = Chroma.from_documents(
                documents=self.documents,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
            
            # Configura retriever
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
        except Exception as e:
            if 'st' in globals():
                st.error(f"Erro ao criar vectorstore de química: {str(e)}")
            else:
                print(f"Erro ao criar vectorstore de química: {str(e)}")
            raise
    
    def load_existing_vectorstore(self) -> bool:
        """Carrega vectorstore existente se disponível"""
        try:
            if os.path.exists(self.persist_directory) and self.embeddings:
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                
                self.retriever = self.vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}
                )
                
                # Testa se a vectorstore tem conteúdo
                try:
                    test_docs = self.retriever.invoke("química")
                    print(f"✅ VectorStore de Química carregada com {len(test_docs)} documentos de teste")
                    
                    # Simula documents para estatísticas
                    sample_docs = self.vectorstore.similarity_search("química", k=100)
                    self.documents = sample_docs
                    print(f"📊 Amostra de química carregada: {len(sample_docs)} chunks")
                    
                except Exception as e:
                    print(f"⚠️ Erro no teste da VectorStore de Química: {e}")
                
                return True
        except Exception as e:
            if 'st' in globals():
                st.warning(f"Não foi possível carregar vectorstore de química existente: {str(e)}")
            else:
                print(f"Não foi possível carregar vectorstore de química existente: {str(e)}")
        
        return False
    
    def create_rag_chain(self, api_key: str):
        """Cria chain RAG conversacional para Química"""
        if not self.retriever:
            raise Exception("Retriever não configurado")
        
        # Configura LLM
        llm = GroqLLM(api_key=api_key)
        
        # Configura memória
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )
        
        # Cria chain conversacional
        self.rag_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=self.retriever,
            memory=self.memory,
            return_source_documents=True,
            verbose=False
        )
        
        # Personaliza o prompt para Química
        self.rag_chain.combine_docs_chain.llm_chain.prompt.template = """
Você é a Professora Luciana, especialista em química do ENEM.

CONTEXTO: {context}
PERGUNTA: {question}

INSTRUÇÕES CRÍTICAS:
🚫 NUNCA mostre seu raciocínio interno
🚫 NUNCA use pensamentos como "Vou analisar...", "Preciso calcular...", etc.
🚫 NUNCA duplique informações
✅ Responda DIRETAMENTE a pergunta
✅ Use fórmulas em LaTeX: $ ou $$
✅ Seja didática para jovem de 17 anos

FORMATO DA RESPOSTA:
1. 🎬 INICIE com analogia das séries da Sther (FRIENDS, Big Bang Theory, Stranger Things, Grey's Anatomy, WandaVision)
2. 👋 Cumprimento: "Olá Sther!"
3. 📚 Explicação DIRETA da química
4. 📝 Exemplo prático quando relevante
5. 🎯 Conecte de volta com a analogia
6. ❓ Termine perguntando sobre exercícios

EXEMPLO DE FÓRMULAS:
- Concentração: $C = \\frac{{n}}{{V}}$
- pH: $$pH = -\\log[H^+]$$
- Equação química: $$2H_2 + O_2 \\rightarrow 2H_2O$$

EXEMPLO DE RESPOSTA CORRETA:
"🎬 Como Monica organizava seus ingredientes na cozinha - na química, cada elemento tem seu lugar na tabela periódica!

Olá Sther! Para calcular concentração molar, usamos:

$$C = \\frac{{n}}{{V}}$$

onde $C$ é concentração (mol/L), $n$ é número de mols e $V$ é volume (L).

**Exemplo:** Para preparar 500 mL de solução 0,1 M de NaCl:
$$n = C \\times V = 0,1 \\times 0,5 = 0,05\\,mol$$

Como Monica sempre sabia exatamente quanto de cada ingrediente usar - na química é igual!

Que tal praticar com alguns exercícios do ENEM sobre este tópico, Sther?"

RESPONDA AGORA SEGUINDO EXATAMENTE ESTE FORMATO:"""
    
    def get_response(self, question: str) -> Dict[str, Any]:
        """Gera resposta usando RAG"""
        if not self.rag_chain:
            return {
                "answer": "Sistema RAG de Química não inicializado. Configure sua API key.",
                "source_documents": []
            }
        
        try:
            result = self.rag_chain({
                "question": question,
                "chat_history": []
            })
            
            # Remove raciocínio interno da resposta
            if "answer" in result:
                answer = result["answer"]
                answer = self._remove_reasoning_from_response(answer)
                
                # Aplica formatação melhorada na resposta
                try:
                    from math_formatter import format_professor_response
                    result["answer"] = format_professor_response(answer)
                except:
                    result["answer"] = answer
            
            return result
            
        except Exception as e:
            return {
                "answer": f"Erro ao gerar resposta: {str(e)}",
                "source_documents": []
            }
    
    def _remove_reasoning_from_response(self, response: str) -> str:
        """Remove qualquer raciocínio interno que ainda apareça na resposta"""
        import re
        
        # Padrões de raciocínio a serem removidos
        reasoning_patterns = [
            r'<think>.*?</think>',  # Tags de pensamento
            r'\*pensa\*.*?\*',      # *pensa* algo *
            r'\*analisa\*.*?\*',    # *analisa* algo *
            r'Vou analisar.*?\.',   # Frases de análise
            r'Preciso calcular.*?\.',  # Frases de cálculo
            r'Vamos pensar.*?\.',   # Frases de pensamento
            r'Deixe-me pensar.*?\.',# Frases de pensamento
            r'Analisando.*?\.',     # Análise
            r'Primeiro.*?vou.*?\.',  # Primeiro vou...
            r'\[Pensamento:.*?\]',  # [Pensamento: ...]
            r'\(pensando.*?\)',     # (pensando...)
        ]
        
        # Remove padrões de raciocínio
        for pattern in reasoning_patterns:
            response = re.sub(pattern, '', response, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove múltiplas quebras de linha
        response = re.sub(r'\n\s*\n\s*\n', '\n\n', response)
        
        # Remove espaços em branco extras
        response = response.strip()
        
        return response
    
    def search_relevant_content(self, query: str, k: int = 3) -> List[Document]:
        """Busca conteúdo relevante"""
        if not self.retriever:
            return []
        
        try:
            docs = self.retriever.get_relevant_documents(query)
            return docs[:k]
        except Exception as e:
            if 'st' in globals():
                st.error(f"Erro na busca de química: {str(e)}")
            else:
                print(f"Erro na busca de química: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas da base de conhecimento de química"""
        stats = {
            "total_documentos": len(self.documents),
            "vectorstore_inicializado": self.vectorstore is not None,
            "retriever_configurado": self.retriever is not None,
            "pasta_quimica": self.chemistry_folder_path,
            "diretorio_persistencia": self.persist_directory
        }
        
        # Conta tópicos únicos
        topics = set()
        for doc in self.documents:
            topic = doc.metadata.get("topic", "Desconhecido")
            topics.add(topic)
        
        stats["topicos_unicos"] = len(topics)
        stats["lista_topicos"] = sorted(list(topics))
        
        return stats
    
    def clear_memory(self):
        """Limpa memória da conversa"""
        if self.memory:
            self.memory.clear()

# Instância global
local_chemistry_rag = LocalChemistryRAG() 