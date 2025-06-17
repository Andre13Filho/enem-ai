"""
Sistema RAG Local para Professor Carlos - Matemática
Processa documentos locais da pasta matemática usando LangChain
"""

import streamlit as st
import os
from typing import Dict, List, Any, Optional
from pathlib import Path
import tempfile
import hashlib

# LangChain imports
from langchain_community.document_loaders import TextLoader, DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain.chains import ConversationalRetrievalChain
try:
    from langchain_community.memory import ConversationBufferMemory
except ImportError:
    from langchain.memory import ConversationBufferMemory
from langchain.llms.base import LLM
from langchain.schema import BaseMessage, AIMessage, HumanMessage
from langchain.callbacks.manager import CallbackManagerForLLMRun

# Document processors
try:
    import docx
    from pypdf import PdfReader
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False

# OpenAI for DeepSeek
from groq import Groq

class GroqLLM(LLM):
    """LLM personalizado para DeepSeek R1 Distill via Groq"""
    
    api_key: str
    model_name: str = "deepseek-r1-distill-llama-70b"
    
    class Config:
        arbitrary_types_allowed = True
    
    def __init__(self, api_key: str, **kwargs):
        super().__init__(api_key=api_key, model_name="deepseek-r1-distill-llama-70b", **kwargs)
        self._client = Groq(api_key=api_key)
    
    @property
    def _llm_type(self) -> str:
        return "groq"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        try:
            response = self._client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=2048
            )
            return response.choices[0].message.content
        except Exception as e:
            from encoding_utils import safe_api_error
            return safe_api_error(e)

class LocalMathRAG:
    """Sistema RAG para documentos de matemática locais"""
    
    def __init__(self, math_folder_path: str = "./matemática"):
        # Configuração adaptativa para cloud/local
        try:
            from cloud_config import get_config
            self.cloud_config = get_config()
            self.math_folder_path = self.cloud_config.get_documents_path("matemática")
            self.persist_directory = self.cloud_config.get_vectorstore_path("math")
        except ImportError:
            # Fallback para configuração local
            self.cloud_config = None
            self.math_folder_path = math_folder_path
            self.persist_directory = "./chroma_math_vectorstore"
            
        self.vectorstore = None
        self.retriever = None
        self.memory = None
        self.rag_chain = None
        self.embeddings = None
        self.documents = []
        self.processed_files = set()
        
        # Inicializa embeddings
        self._setup_embeddings()
    
    def _setup_embeddings(self):
        """Configura embeddings usando HuggingFace"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception as e:
            st.error(f"Erro ao configurar embeddings: {str(e)}")
    
    def process_math_documents(self) -> bool:
        """Processa todos os documentos da pasta matemática"""
        if not os.path.exists(self.math_folder_path):
            st.error(f"Pasta não encontrada: {self.math_folder_path}")
            return False
        
        if not DOCUMENT_PROCESSING_AVAILABLE:
            st.error("Bibliotecas de processamento não disponíveis. Instale: pip install python-docx pypdf")
            return False
        
        try:
            # Lista todos os arquivos na pasta matemática
            math_files = []
            for root, dirs, files in os.walk(self.math_folder_path):
                for file in files:
                    if file.lower().endswith(('.docx', '.pdf', '.txt')):
                        math_files.append(os.path.join(root, file))
            
            st.info(f"Encontrados {len(math_files)} arquivos para processar...")
            print(f"🔍 Debug: Arquivos encontrados: {math_files}")
            
            # Processa cada arquivo
            all_documents = []
            progress_bar = st.progress(0)
            
            for i, file_path in enumerate(math_files):
                try:
                    print(f"🔍 Debug: Processando arquivo {file_path}")
                    docs = self._process_single_file(file_path)
                    all_documents.extend(docs)
                    progress_bar.progress((i + 1) / len(math_files))
                    print(f"✅ Debug: Arquivo {file_path} processado com {len(docs)} chunks")
                except Exception as e:
                    print(f"❌ Debug: Erro ao processar {file_path}: {str(e)}")
                    st.warning(f"Erro ao processar {file_path}: {str(e)}")
                    continue
            
            self.documents = all_documents
            
            if not self.documents:
                st.error("Nenhum documento foi processado com sucesso")
                return False
            
            # Cria ou atualiza vectorstore
            print(f"🔍 Debug: Criando vectorstore com {len(self.documents)} documentos")
            self._create_vectorstore()
            
            st.success(f"✅ Processados {len(self.documents)} chunks de {len(math_files)} arquivos")
            return True
            
        except Exception as e:
            print(f"❌ Debug: Erro geral no processamento: {str(e)}")
            st.error(f"Erro no processamento: {str(e)}")
            return False
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Processa um único arquivo"""
        file_extension = Path(file_path).suffix.lower()
        file_name = Path(file_path).name
        
        # Extrai conteúdo baseado no tipo
        if file_extension == '.docx':
            content = self._extract_docx_content(file_path)
        elif file_extension == '.pdf':
            content = self._extract_pdf_content(file_path)
        elif file_extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                # Tenta outros encodings comuns
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    st.warning(f"Não foi possível decodificar o arquivo {file_path}")
                    return []
        else:
            return []
        
        if not content.strip():
            return []
        
        # Cria documento LangChain
        doc = Document(
            page_content=content,
            metadata={
                "source": file_name,
                "file_path": file_path,
                "type": "matematica",
                "topic": self._extract_topic_from_filename(file_name)
            }
        )
        
        # Divide em chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        chunks = text_splitter.split_documents([doc])
        return chunks
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivos DOCX"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Garante encoding UTF-8
                    text = paragraph.text.strip()
                    if isinstance(text, str):
                        text = text.encode('utf-8', errors='replace').decode('utf-8')
                    content_parts.append(text)
            
            return "\n\n".join(content_parts)
        except Exception as e:
            st.warning(f"Erro ao processar DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivos PDF"""
        try:
            reader = PdfReader(file_path)
            content_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    # Garante encoding UTF-8
                    text = text.strip()
                    if isinstance(text, str):
                        text = text.encode('utf-8', errors='replace').decode('utf-8')
                    content_parts.append(text)
            
            return "\n\n".join(content_parts)
        except Exception as e:
            st.warning(f"Erro ao processar PDF {file_path}: {str(e)}")
            return ""
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extrai tópico do nome do arquivo"""
        # Remove extensão
        topic = Path(filename).stem
        
        # Limpa caracteres especiais
        topic = topic.replace("_", " ").replace("-", " ")
        
        # Capitaliza primeira letra
        return topic.title()
    
    def _create_vectorstore(self):
        """Cria vectorstore com fallback para ambiente cloud"""
        if not self.documents:
            raise ValueError("Nenhum documento disponível para criar vectorstore")
        
        try:
            # Primeira tentativa: vectorstore persistente
            try:
                os.makedirs(self.persist_directory, exist_ok=True)
                self.vectorstore = Chroma.from_documents(
                    documents=self.documents,
                    embedding=self.embeddings,
                    persist_directory=self.persist_directory
                )
                self.vectorstore.persist()
                print(f"✅ Vectorstore persistente criado em {self.persist_directory}")
            except Exception as persist_error:
                print(f"⚠️ Falha na criação persistente: {persist_error}")
                # Segunda tentativa: vectorstore em memória
                self.vectorstore = Chroma.from_documents(
                    documents=self.documents,
                    embedding=self.embeddings
                )
                print("✅ Vectorstore em memória criado com sucesso")
            
            # Configura retriever
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
            # Teste básico de funcionamento
            test_docs = self.retriever.invoke("matemática")
            print(f"✅ Vectorstore funcionando - teste retornou {len(test_docs)} documentos")
            
        except Exception as e:
            print(f"❌ Erro crítico na criação do vectorstore: {str(e)}")
            raise
    
    def load_existing_vectorstore(self) -> bool:
        """Carrega vectorstore existente ou cria um novo em memória"""
        try:
            # Primeiro tenta carregar vectorstore persistente
            if os.path.exists(self.persist_directory):
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                self.retriever = self.vectorstore.as_retriever(search_kwargs={"k": 5})
                print("✅ Vectorstore persistente carregado com sucesso")
                return True
        except Exception as e:
            print(f"⚠️ Erro ao carregar vectorstore persistente: {str(e)}")
        
        # Se falhar, tenta criar vectorstore em memória
        try:
            # Processa documentos se ainda não foram processados
            if not self.documents:
                print("🔄 Processando documentos para vectorstore em memória...")
                if not self.process_math_documents():
                    return False
            
            if self.documents:
                # Cria vectorstore em memória (sem persistência)
                self.vectorstore = Chroma.from_documents(
                    documents=self.documents,
                    embedding=self.embeddings
                )
                self.retriever = self.vectorstore.as_retriever(search_kwargs={"k": 5})
                print("✅ Vectorstore em memória criado com sucesso")
                return True
            else:
                print("❌ Nenhum documento disponível para criar vectorstore")
                return False
                
        except Exception as e:
            print(f"❌ Erro ao criar vectorstore em memória: {str(e)}")
            return False
    
    def create_rag_chain(self, api_key: str):
        """Cria chain RAG conversacional"""
        if not self.retriever:
            raise Exception("Retriever não configurado")
        
        try:
            print(f"🔗 Criando chain RAG com API key: {api_key[:10]}...")
            
            # Configura LLM
            llm = GroqLLM(api_key=api_key)
            print("✅ LLM GroqLLM configurado")
            
            # Configura memória
            self.memory = ConversationBufferMemory(
                memory_key="chat_history",
                return_messages=True,
                output_key="answer"
            )
            print("✅ Memória conversacional configurada")
            
            # Cria chain conversacional
            self.rag_chain = ConversationalRetrievalChain.from_llm(
                llm=llm,
                retriever=self.retriever,
                memory=self.memory,
                return_source_documents=True,
                verbose=False
            )
            print("✅ Chain RAG conversacional criada")
            
        except Exception as e:
            print(f"❌ Erro ao criar chain RAG: {str(e)}")
            print(f"• Tipo do erro: {type(e).__name__}")
            print(f"• API key fornecida: {api_key[:10] if api_key else 'None'}...")
            print(f"• Retriever disponível: {self.retriever is not None}")
            raise Exception(f"Falha na criação da chain RAG: {str(e)}")
        
        # Personaliza o prompt
        if self.rag_chain:
            try:
                self.rag_chain.combine_docs_chain.llm_chain.prompt.template = """
Você é o Professor Carlos, especialista em matemática do ENEM. Responda APENAS como professor, SEM mostrar seu raciocínio interno.

CONTEXTO: {context}
PERGUNTA: {question}

REGRAS OBRIGATÓRIAS:
1. NÃO mostre seu processo de pensamento
2. NÃO duplique informações
3. NÃO use colchetes [ ] para fórmulas
4. SEMPRE use LaTeX entre $ ou $$
5. Seja direto e claro

FORMATAÇÃO DE FÓRMULAS (OBRIGATÓRIA):
- Inline: $f(x) = ax^2 + bx + c$
- Destaque: $$x = \\frac{{-b \\pm \\sqrt{{b^2-4ac}}}}{{2a}}$$
- Frações: $$\\frac{{numerador}}{{denominador}}$$
- Expoentes: $x^2$, $y^3$
- Raízes: $$\\sqrt{{a^2 + b^2}}$$

EXEMPLO CORRETO:
"Para polígonos, usamos a fórmula:

$$D = \\frac{{n(n-3)}}{{2}}$$

onde $n$ é o número de lados.

**Exemplo prático:** Para um pentágono ($n = 5$):
$$D = \\frac{{5(5-3)}}{{2}} = \\frac{{5 \\times 2}}{{2}} = 5$$"

Responda de forma didática para estudante de 17 anos, usando o contexto fornecido:"""
                print("✅ Prompt personalizado configurado")
            except Exception as prompt_error:
                print(f"⚠️ Erro ao configurar prompt: {str(prompt_error)}")
        
        # Verificação final
        if self.rag_chain is None:
            raise Exception("Chain RAG é None após criação - falha crítica")
        
        print("🎉 Chain RAG criada e configurada com sucesso!")
    
    def get_response(self, question: str) -> Dict[str, Any]:
        """Gera resposta usando RAG"""
        if not self.rag_chain:
            print("❌ get_response: rag_chain é None")
            print(f"• Retriever disponível: {self.retriever is not None}")
            print(f"• Documentos carregados: {len(self.documents) if self.documents else 0}")
            print(f"• Vectorstore configurado: {self.vectorstore is not None}")
            
            return {
                "answer": "⚠️ Sistema RAG não inicializado. A cadeia RAG não foi criada. Verifique: 1) API key válida, 2) Documentos processados, 3) Vectorstore criado.",
                "source_documents": []
            }
        
        try:
            print(f"🤖 Gerando resposta para: {question[:50]}...")
            result = self.rag_chain({
                "question": question,
                "chat_history": []
            })
            print("✅ Resposta gerada com sucesso")
            
            # Aplica formatação melhorada na resposta
            try:
                from math_formatter import format_professor_response
                if "answer" in result:
                    result["answer"] = format_professor_response(result["answer"])
            except ImportError:
                print("⚠️ math_formatter não disponível - pulando formatação")
            
            return result
            
        except Exception as e:
            error_msg = str(e)
            print(f"❌ Erro ao gerar resposta: {error_msg}")
            
            # Verifica se é erro de API key
            if "401" in error_msg or "invalid" in error_msg.lower() or "unauthorized" in error_msg.lower():
                return {
                    "answer": "🔑 Erro de autenticação com a API Groq. Verifique se sua API key está correta e válida.",
                    "source_documents": []
                }
            
            return {
                "answer": f"❌ Erro ao gerar resposta: {error_msg}",
                "source_documents": []
            }
    
    def search_relevant_content(self, query: str, k: int = 3) -> List[Document]:
        """Busca conteúdo relevante"""
        if not self.retriever:
            return []
        
        try:
            docs = self.retriever.get_relevant_documents(query)
            return docs[:k]
        except Exception as e:
            st.error(f"Erro na busca: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas da base de conhecimento"""
        stats = {
            "total_documentos": len(self.documents),
            "vectorstore_inicializado": self.vectorstore is not None,
            "retriever_configurado": self.retriever is not None,
            "pasta_matematica": self.math_folder_path,
            "diretorio_persistencia": self.persist_directory
        }
        
        # Conta tópicos únicos
        topics = set()
        for doc in self.documents:
            topic = doc.metadata.get("topic", "Desconhecido")
            topics.add(topic)
        
        stats["topicos_unicos"] = len(topics)
        stats["lista_topicos"] = sorted(list(topics))
        
        return stats
    
    def clear_memory(self):
        """Limpa memória da conversa"""
        if self.memory:
            self.memory.clear()

# Instância global
local_math_rag = LocalMathRAG() 