#!/usr/bin/env python3
"""
Sistema RAG Local para História
Professor Eduardo - ENEM AI Helper

Baseado no sistema dos outros professores, adaptado para história
"""

import os
import streamlit as st
from pathlib import Path
from typing import List, Dict, Any, Optional
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain.llms.base import LLM
from groq import Groq
import docx
import PyPDF2
import re

class GroqLLM(LLM):
    """LLM personalizado para DeepSeek R1 Distill via Groq"""
    
    api_key: str
    model_name: str = "deepseek-r1-distill-llama-70b"
    temperature: float = 0.1
    max_tokens: int = 4000
    client: Optional[Any] = None
    
    class Config:
        arbitrary_types_allowed = True
    
    def __init__(self, api_key: str, **kwargs):
        super().__init__(api_key=api_key, **kwargs)
        self.client = Groq(api_key=api_key)
    
    @property
    def _llm_type(self) -> str:
        return "groq"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        """Chama a API Groq"""
        try:
            # Adiciona instruções específicas contra mostrar raciocínio
            system_message = """Você é um professor. NUNCA mostre seu raciocínio interno ou processo de pensamento. Responda DIRETAMENTE como um professor explicando para um aluno. NÃO use frases como "Vou analisar", "Preciso calcular", "Vamos pensar", etc."""
            
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                stop=stop
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Erro na API: {str(e)}"

class LocalHistoryRAG:
    """Sistema RAG para História com vectorstore local"""
    
    def __init__(self, history_folder_path: str = "./historia"):
        # Configuração adaptativa para cloud/local
        try:
            from cloud_config import get_config
            self.cloud_config = get_config()
            self.history_folder_path = self.cloud_config.get_documents_path("história")
            self.persist_directory = self.cloud_config.get_vectorstore_path("history")
        except ImportError:
            # Fallback para configuração local
            self.cloud_config = None
            self.history_folder_path = history_folder_path
            self.persist_directory = "./chroma_history_db"
        self.documents = []
        self.vectorstore = None
        self.retriever = None
        self.rag_chain = None
        self.memory = None
        self.embeddings = None
        
        # Configura embeddings
        self._setup_embeddings()
        
        # Tenta carregar vectorstore existente
        if not self.load_existing_vectorstore():
            # Se não existir, processa documentos se a pasta existir
            if os.path.exists(self.history_folder_path):
                self.process_history_documents()
    
    def _setup_embeddings(self):
        """Configura embeddings multilíngues"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception as e:
            st.error(f"Erro ao configurar embeddings: {str(e)}")
            self.embeddings = None
    
    def process_history_documents(self) -> bool:
        """Processa documentos de história e cria vectorstore"""
        try:
            if not os.path.exists(self.history_folder_path):
                st.warning(f"Pasta {self.history_folder_path} não encontrada")
                return False
            
            history_files = list(Path(self.history_folder_path).glob("*"))
            
            if not history_files:
                st.warning(f"Nenhum arquivo encontrado em {self.history_folder_path}")
                return False
            
            # Processa cada arquivo
            for file_path in history_files:
                if file_path.is_file():
                    docs = self._process_single_file(str(file_path))
                    self.documents.extend(docs)
            
            if not self.documents:
                st.warning("Nenhum documento processado")
                return False
            
            # Cria vectorstore
            self._create_vectorstore()
            
            st.success(f"✅ Processados {len(self.documents)} documentos de história")
            return True
            
        except Exception as e:
            st.error(f"Erro ao processar documentos de história: {str(e)}")
            return False
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Processa um único arquivo"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            # Determina método de extração baseado na extensão
            if file_extension == '.txt':
                content = self._extract_txt_content(file_path)
            elif file_extension == '.docx':
                content = self._extract_docx_content(file_path)
            elif file_extension == '.pdf':
                content = self._extract_pdf_content(file_path)
            else:
                return []
            
            if not content or len(content.strip()) < 50:
                return []
            
            # Cria divisor de texto
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", ".", " ", ""]
            )
            
            # Divide texto em chunks
            chunks = text_splitter.split_text(content)
            
            # Cria documentos
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "filename": Path(file_path).name,
                        "topic": self._extract_topic_from_filename(Path(file_path).name),
                        "chunk_id": i,
                        "subject": "História"
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            st.warning(f"Erro ao processar {file_path}: {str(e)}")
            return []
    
    def _extract_txt_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .txt"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    if isinstance(content, str):
                        content = content.encode('utf-8', errors='replace').decode('utf-8')
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                st.warning(f"Erro ao ler arquivo {file_path}: {str(e)}")
                return ""
        
        st.warning(f"Não foi possível decodificar {file_path}")
        return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .docx"""
        try:
            doc = docx.Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            return '\n'.join(content)
        except Exception as e:
            st.warning(f"Erro ao ler DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .pdf"""
        try:
            content_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Verifica se o PDF está criptografado
                if pdf_reader.is_encrypted:
                    # Tenta descriptografar com senha vazia
                    try:
                        pdf_reader.decrypt("")
                    except:
                        st.warning(f"PDF {Path(file_path).name} está protegido por senha - pulando...")
                        return ""
                
                # Extrai texto das páginas
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if isinstance(text, str):
                            text = text.encode('utf-8', errors='replace').decode('utf-8')
                        if text.strip():
                            content_parts.append(text)
                    except Exception as page_error:
                        st.warning(f"Erro na página {i+1} do PDF {Path(file_path).name}: {str(page_error)}")
                        continue
            
            if content_parts:
                return "\n\n".join(content_parts)
            else:
                st.info(f"PDF {Path(file_path).name} processado mas sem texto extraído")
                return ""
                
        except Exception as e:
            error_msg = str(e)
            if "PyCryptodome" in error_msg:
                st.error(f"PDF {Path(file_path).name} requer descriptografia avançada - instale: pip install PyCryptodome")
            else:
                st.warning(f"Erro ao processar PDF {Path(file_path).name}: {error_msg}")
            return ""
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extrai tópico do nome do arquivo"""
        # Remove extensão
        topic = Path(filename).stem
        
        # Limpa caracteres especiais
        topic = topic.replace("_", " ").replace("-", " ")
        
        # Capitaliza primeira letra
        return topic.title()
    
    def _create_vectorstore(self):
        """Cria ou atualiza o vectorstore"""
        try:
            if not self.embeddings:
                raise Exception("Embeddings não configurados")
            
            # Cria vectorstore ChromaDB
            self.vectorstore = Chroma.from_documents(
                documents=self.documents,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
            
            # Configura retriever
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
        except Exception as e:
            st.error(f"Erro ao criar vectorstore de história: {str(e)}")
            raise
    
    def load_existing_vectorstore(self) -> bool:
        """Carrega vectorstore existente se disponível"""
        try:
            if os.path.exists(self.persist_directory) and self.embeddings:
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                
                self.retriever = self.vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}
                )
                
                # Testa se a vectorstore tem conteúdo
                try:
                    test_docs = self.retriever.invoke("história")
                    print(f"✅ VectorStore de História carregada com {len(test_docs)} documentos de teste")
                    
                    # Simula documents para estatísticas
                    sample_docs = self.vectorstore.similarity_search("história", k=100)
                    self.documents = sample_docs
                    print(f"📊 Amostra de história carregada: {len(sample_docs)} chunks")
                    
                except Exception as e:
                    print(f"⚠️ Erro no teste da VectorStore de História: {e}")
                
                return True
        except Exception as e:
            if 'st' in globals():
                st.warning(f"Não foi possível carregar vectorstore de história existente: {str(e)}")
            else:
                print(f"Não foi possível carregar vectorstore de história existente: {str(e)}")
        
        return False
    
    def create_rag_chain(self, api_key: str):
        """Cria chain RAG conversacional para História"""
        if not self.retriever:
            raise Exception("Retriever não configurado")
        
        # Configura LLM
        llm = GroqLLM(api_key=api_key)
        
        # Configura memória
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )
        
        # Cria chain conversacional
        self.rag_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=self.retriever,
            memory=self.memory,
            return_source_documents=True,
            verbose=False
        )
        
        # Personaliza o prompt para História
        self.rag_chain.combine_docs_chain.llm_chain.prompt.template = """
Você é o Professor Eduardo, especialista em história do ENEM.

CONTEXTO: {context}
PERGUNTA: {question}

INSTRUÇÕES CRÍTICAS:
🚫 NUNCA mostre seu raciocínio interno
🚫 NUNCA use pensamentos como "Vou analisar...", "Preciso explicar...", etc.
🚫 NUNCA duplique informações
✅ Responda DIRETAMENTE a pergunta
✅ Use datas e contextos históricos
✅ Seja didático para jovem de 17 anos

FORMATO DA RESPOSTA:
1. 🎬 INICIE com analogia das séries da Sther (FRIENDS, Big Bang Theory, Stranger Things, Grey's Anatomy, WandaVision)
2. 👋 Cumprimento: "Olá Sther!"
3. 📚 Explicação DIRETA da história
4. 📝 Exemplo prático com datas e contextos
5. 🎯 Conecte de volta com a analogia
6. ❓ Termine perguntando sobre exercícios

## 🎭 ANALOGIAS DAS SÉRIES POR TÓPICO (USE SEMPRE):

### 🍕 FRIENDS:
- **História Antiga**: "Como Monica organizava suas lembranças - cada civilização antiga tinha sua organização específica!"
- **Revoluções**: "Pense nas revoluções como quando Ross e Rachel se separavam - mudanças que transformavam tudo!"
- **Períodos Históricos**: "Como quando o grupo se reunia no Central Perk - cada época tinha seu 'ponto de encontro' cultural!"
- **Guerras**: "Lembra quando Chandler explicava conflitos? 'Could this BE more histórico?'"

### 🧪 THE BIG BANG THEORY:
- **Civilizações**: "Como Sheldon explicava: 'Bazinga! A história é simples quando você entende as causas e consequências!'"
- **Cronologia**: "Como Leonard organizava experimentos - cada evento histórico tem sua sequência lógica!"
- **Revoluções**: "Como Howard calculava trajetórias - as revoluções seguem padrões previsíveis!"
- **Impérios**: "Raj usaria a história para explicar como impérios nascem e caem em ciclos!"

### 🌌 STRANGER THINGS:
- **Idade Média**: "Como Dustin explicava sobre o mundo medieval - cheio de mistérios e hierarquias!"
- **Colonização**: "Como eles exploravam territórios desconhecidos - igual aos grandes navegadores!"
- **Revoluções**: "Como quando Will enfrentava mudanças no Mundo Invertido - as revoluções transformam sociedades!"

### 🏥 GREY'S ANATOMY:
- **História do Brasil**: "Como Meredith dizia: 'Você é minha pessoa!' - cada período histórico tem suas características únicas!"
- **Política**: "Como analisavam casos complexos no Seattle Grace - a política envolve múltiplas variáveis!"
- **Sociedade**: "Derek sempre conectava sintomas - na história, eventos estão sempre interligados!"

### ✨ WANDAVISION:
- **Transformações Históricas**: "Como Wanda criava realidades diferentes - cada época histórica teve sua 'realidade' social!"
- **Revoluções**: "Visão calculava mudanças com precisão - as revoluções seguem padrões históricos!"
- **Evolução Social**: "Como Wanda evoluía seus poderes - as sociedades evoluem através dos tempos!"

EXEMPLO DE RESPOSTA CORRETA:
"🎬 Como Monica organizava seus álbuns de fotos por épocas - na história, organizamos eventos por períodos!

Olá Sther! A Revolução Francesa (1789-1799) foi um marco que transformou a Europa:

**Causas principais:**
- Crise econômica e social
- Influência do Iluminismo
- Desigualdade entre os Estados

**Fases importantes:**
- 1789: Queda da Bastilha
- 1792: Proclamação da República
- 1799: Golpe de Napoleão

Como Monica sempre sabia onde encontrar cada lembrança - entender a cronologia nos ajuda a organizar a história!

Que tal praticar com alguns exercícios do ENEM sobre este tópico, Sther?"

RESPONDA AGORA SEGUINDO EXATAMENTE ESTE FORMATO:"""
    
    def get_response(self, question: str) -> Dict[str, Any]:
        """Gera resposta usando RAG"""
        if not self.rag_chain:
            return {
                "answer": "Sistema RAG de História não inicializado. Configure sua API key.",
                "source_documents": []
            }
        
        try:
            result = self.rag_chain({
                "question": question,
                "chat_history": []
            })
            
            # Remove raciocínio interno da resposta
            if "answer" in result:
                answer = result["answer"]
                answer = self._remove_reasoning_from_response(answer)
                result["answer"] = answer
            
            return result
            
        except Exception as e:
            return {
                "answer": f"Erro ao gerar resposta: {str(e)}",
                "source_documents": []
            }
    
    def _remove_reasoning_from_response(self, response: str) -> str:
        """Remove qualquer raciocínio interno que ainda apareça na resposta"""
        import re
        
        # Padrões de raciocínio a serem removidos
        reasoning_patterns = [
            r'<think>.*?</think>',  # Tags de pensamento
            r'\*pensa\*.*?\*',      # *pensa* algo *
            r'\*analisa\*.*?\*',    # *analisa* algo *
            r'Vou analisar.*?\.',   # Frases de análise
            r'Preciso.*?\.',        # Frases de necessidade
            r'Deixe-me.*?\.',       # Frases de reflexão
            r'Para responder.*?\.',  # Frases de preparação
            r'Baseado.*?\.',        # Frases de base
        ]
        
        cleaned_response = response
        for pattern in reasoning_patterns:
            cleaned_response = re.sub(pattern, '', cleaned_response, flags=re.DOTALL)
        
        # Remove linhas vazias múltiplas
        cleaned_response = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_response)
        
        return cleaned_response.strip()
    
    def search_relevant_content(self, query: str, k: int = 3) -> List[Document]:
        """Busca conteúdo relevante no vectorstore"""
        if not self.vectorstore:
            return []
        
        try:
            return self.vectorstore.similarity_search(query, k=k)
        except Exception as e:
            st.error(f"Erro na busca: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas do sistema"""
        return {
            "total_documents": len(self.documents),
            "vectorstore_exists": self.vectorstore is not None,
            "rag_chain_ready": self.rag_chain is not None,
            "history_folder_exists": os.path.exists(self.history_folder_path),
            "persist_directory": self.persist_directory
        }
    
    def clear_memory(self):
        """Limpa memória de conversação"""
        if self.memory:
            self.memory.clear()

# Instância global do sistema RAG
history_rag_system = LocalHistoryRAG() 