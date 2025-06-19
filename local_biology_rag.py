#!/usr/bin/env python3
"""
Sistema RAG Local para Biologia
Professor Roberto - ENEM AI Helper

Baseado no sistema de física, adaptado para biologia
"""

import os
import streamlit as st
from pathlib import Path
from typing import List, Dict, Any, Optional
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain.llms.base import LLM
from groq import Groq
import docx
import PyPDF2
import re

class GroqLLM(LLM):
    """LLM personalizado para DeepSeek R1 Distill via Groq"""
    
    api_key: str
    model_name: str = "deepseek-r1-distill-llama-70b"
    temperature: float = 0.1
    max_tokens: int = 4000
    
    class Config:
        arbitrary_types_allowed = True
    
    def __init__(self, api_key: str, **kwargs):
        super().__init__(api_key=api_key, **kwargs)
    
    @property
    def _llm_type(self) -> str:
        return "groq"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        """Chama a API Groq"""
        try:
            # Adiciona instruções específicas contra mostrar raciocínio
            system_message = """Você é um professor. NUNCA mostre seu raciocínio interno ou processo de pensamento. Responda DIRETAMENTE como um professor explicando para um aluno. NÃO use frases como "Vou analisar", "Preciso calcular", "Vamos pensar", etc."""
            
            # Cria uma nova instância do cliente a cada chamada para evitar cache corrompido
            client = Groq(api_key=self.api_key)
            response = client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                stop=stop
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Erro na API: {str(e)}"

class LocalBiologyRAG:
    """Sistema RAG para Biologia com vectorstore local"""
    
    def __init__(self, biology_folder_path: str = "./biologia"):
        # Configuração adaptativa para cloud/local
        try:
            from cloud_config import get_config
            self.cloud_config = get_config()
            self.biology_folder_path = self.cloud_config.get_documents_path("biologia")
            self.persist_directory = self.cloud_config.get_vectorstore_path("biology")
        except ImportError:
            # Fallback para configuração local
            self.cloud_config = None
            self.biology_folder_path = biology_folder_path
            self.persist_directory = "./chroma_biology_db"
            
        self.documents = []
        self.vectorstore = None
        self.retriever = None
        self.rag_chain = None
        self.memory = None
        self.embeddings = None
        
        # Configura embeddings
        self._setup_embeddings()
        
        # Tenta carregar vectorstore existente
        if not self.load_existing_vectorstore():
            # Se não existir, processa documentos se a pasta existir
            if os.path.exists(self.biology_folder_path):
                self.process_biology_documents()
    
    def _setup_embeddings(self):
        """Configura embeddings multilíngues"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception as e:
            st.error(f"Erro ao configurar embeddings: {str(e)}")
            self.embeddings = None
    
    def process_biology_documents(self) -> bool:
        """Processa documentos de biologia e cria vectorstore"""
        try:
            if not os.path.exists(self.biology_folder_path):
                st.warning(f"Pasta {self.biology_folder_path} não encontrada")
                return False
            
            biology_files = list(Path(self.biology_folder_path).glob("*"))
            
            if not biology_files:
                st.warning(f"Nenhum arquivo encontrado em {self.biology_folder_path}")
                return False
            
            # Processa cada arquivo
            for file_path in biology_files:
                if file_path.is_file():
                    docs = self._process_single_file(str(file_path))
                    self.documents.extend(docs)
            
            if not self.documents:
                st.warning("Nenhum documento processado")
                return False
            
            # Cria vectorstore
            self._create_vectorstore()
            
            st.success(f"✅ Processados {len(self.documents)} documentos de biologia")
            return True
            
        except Exception as e:
            st.error(f"Erro ao processar documentos de biologia: {str(e)}")
            return False
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Processa um único arquivo"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            # Determina método de extração baseado na extensão
            if file_extension == '.txt':
                content = self._extract_txt_content(file_path)
            elif file_extension == '.docx':
                content = self._extract_docx_content(file_path)
            elif file_extension == '.pdf':
                content = self._extract_pdf_content(file_path)
            else:
                return []
            
            if not content or len(content.strip()) < 50:
                return []
            
            # Cria divisor de texto
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", ".", " ", ""]
            )
            
            # Divide texto em chunks
            chunks = text_splitter.split_text(content)
            
            # Cria documentos
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "filename": Path(file_path).name,
                        "topic": self._extract_topic_from_filename(Path(file_path).name),
                        "chunk_id": i,
                        "subject": "Biologia"
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            st.warning(f"Erro ao processar {file_path}: {str(e)}")
            return []
    
    def _extract_txt_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .txt"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    if isinstance(content, str):
                        content = content.encode('utf-8', errors='replace').decode('utf-8')
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                st.warning(f"Erro ao ler arquivo {file_path}: {str(e)}")
                return ""
        
        st.warning(f"Não foi possível decodificar {file_path}")
        return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .docx"""
        try:
            doc = docx.Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            return '\n'.join(content)
        except Exception as e:
            st.warning(f"Erro ao ler DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .pdf"""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Verifica se o PDF está criptografado
                if pdf_reader.is_encrypted:
                    # Tenta descriptografar com senha vazia
                    try:
                        pdf_reader.decrypt("")
                    except:
                        st.warning(f"PDF {Path(file_path).name} está protegido por senha - pulando...")
                        return ""
                
                # Extrai texto das páginas
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content.append(page_text)
                    except Exception as page_error:
                        st.warning(f"Erro na página {i+1} do PDF {Path(file_path).name}: {str(page_error)}")
                        continue
                
            if content:
                return '\n'.join(content)
            else:
                st.info(f"PDF {Path(file_path).name} processado mas sem texto extraído")
                return ""
                
        except Exception as e:
            error_msg = str(e)
            if "PyCryptodome" in error_msg:
                st.error(f"PDF {Path(file_path).name} requer descriptografia avançada - instale: pip install PyCryptodome")
            else:
                st.warning(f"Erro ao ler PDF {Path(file_path).name}: {error_msg}")
            return ""
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extrai tópico do nome do arquivo"""
        # Remove extensão e normaliza
        topic = Path(filename).stem
        return topic.replace('_', ' ').replace('-', ' ').title()
    
    def _create_vectorstore(self):
        """Cria vectorstore a partir dos documentos processados"""
        try:
            if not self.embeddings:
                st.error("Embeddings não configurados")
                return
            
            if not self.documents:
                st.warning("Nenhum documento para criar vectorstore")
                return
            
            # Cria vectorstore com persistência
            self.vectorstore = Chroma.from_documents(
                documents=self.documents,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
            
            # Salva para uso futuro
            self.vectorstore.persist()
            
            st.success("✅ Vectorstore criado com sucesso")
            
        except Exception as e:
            st.error(f"Erro ao criar vectorstore: {str(e)}")
    
    def load_existing_vectorstore(self) -> bool:
        """Carrega vectorstore existente se disponível"""
        try:
            if not os.path.exists(self.persist_directory):
                return False
            
            if not self.embeddings:
                self._setup_embeddings()
                if not self.embeddings:
                    return False
            
            # Carrega vectorstore existente
            self.vectorstore = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings
            )
            
            # Testa se está funcionando
            test_docs = self.vectorstore.similarity_search("teste", k=1)
            
            st.success("✅ Vectorstore de biologia carregado")
            return True
            
        except Exception as e:
            st.warning(f"Não foi possível carregar vectorstore existente: {str(e)}")
            # Remove diretório corrompido
            try:
                import shutil
                shutil.rmtree(self.persist_directory)
            except:
                pass
            return False
    
    def create_rag_chain(self, api_key: str):
        """Cria chain RAG para conversação"""
        try:
            if not self.vectorstore:
                st.error("Vectorstore não disponível")
                return False
            
            # Configura LLM
            llm = GroqLLM(api_key=api_key)
            
            # Configura retriever
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 3}
            )
            
            # Configura memória
            self.memory = ConversationBufferMemory(
                memory_key="chat_history",
                return_messages=True,
                output_key="answer"
            )
            
            # Template específico para biologia
            biology_template = """# IDENTIDADE DO PROFESSOR
Você é o Professor Roberto, especialista em Biologia para o ENEM.
Você é professor particular exclusivo da Sther Souza, uma estudante de 17 anos.

# DIRETRIZES DE COMPORTAMENTO
## Personalidade:
- Didático, paciente e bem-humorado
- Linguagem adequada para jovens de 17 anos
- Seja um professor que gosta de ensinar
- Seja carinhoso, mas profissional
- Seja engraçado, mas profissional
- Dê uma aula para uma jovem de 17 anos que tem dificuldades em biologia

## Metodologia de Ensino:
- Use analogias e exemplos práticos
- Relacione com situações do cotidiano brasileiro
- Foque nos tópicos que caem no ENEM
- Use conceitos e processos biológicos quando aplicável
- Responda sempre em português brasileiro

## Estrutura da Resposta:
1. 🎬 **SEMPRE inicie com uma analogia das séries favoritas da Sther** (FRIENDS, The Big Bang Theory, Stranger Things, Grey's Anatomy, WandaVision)
2. Responda a pergunta de forma clara e completa
3. Use formatação em markdown para organizar o conteúdo
4. Inclua exemplos práticos quando relevante
5. SEMPRE termine perguntando sobre exercícios

## 🎭 ANALOGIAS DAS SÉRIES POR TÓPICO (USE SEMPRE):

### 🍕 FRIENDS:
- **Células**: "Como Monica organizava seu apartamento por funções - na célula, cada organela tem sua função específica!"
- **Genética**: "Pense na genética como o relacionamento de Ross e Rachel - características dos pais são herdadas pelos filhos!"
- **Ecossistemas**: "Como quando o grupo se reunia no Central Perk - cada um tinha seu papel no 'ecossistema' da turma!"
- **Metabolismo**: "Lembra quando Chandler explicava seu trabalho? 'Could this BE more metabólico?'"

### 🧪 THE BIG BANG THEORY:
- **Biologia Molecular**: "Como Sheldon explicava: 'Bazinga! A biologia molecular é simples quando você entende as estruturas!'"
- **Evolução**: "Como Leonard tentava explicar conceitos para Penny - cada espécie tem sua adaptação!"
- **Genética**: "Como Howard calculava probabilidades - a genética é estatística da vida!"
- **Anatomia**: "Raj usaria anatomia para explicar como o corpo funciona em harmonia!"

### 🌌 STRANGER THINGS:
- **Sistema Nervoso**: "Como Dustin explicava sobre as conexões neurais no cérebro!"
- **Microbiologia**: "Como eles investigavam criaturas microscópicas no laboratório de Hawkins!"
- **Genética**: "Como quando Will precisava entender as mutações do Mundo Invertido!"

### 🏥 GREY'S ANATOMY:
- **Anatomia**: "Como Meredith dizia: 'Você é minha pessoa!' - cada órgão tem sua função vital!"
- **Fisiologia**: "Como analisavam o funcionamento do corpo no Seattle Grace!"
- **Patologia**: "Derek sempre analisava sintomas para diagnosticar doenças!"

### ✨ WANDAVISION:
- **Biologia Celular**: "Como Wanda criava realidades celulares perfeitas em Westview!"
- **Genética**: "Visão calculava hereditariedade com precisão para entender mutações!"
- **Evolução**: "Como Wanda evoluía suas habilidades com sua Magia do Caos!"

Use as informações do contexto fornecido e o histórico da conversa para responder de forma precisa e educativa.

CONTEXTO: {context}

PERGUNTA: {question}

HISTÓRICO: {chat_history}

RESPOSTA:"""
            
            # Cria chain conversacional
            self.rag_chain = ConversationalRetrievalChain.from_llm(
                llm=llm,
                retriever=self.retriever,
                memory=self.memory,
                return_source_documents=True,
                verbose=False,
                combine_docs_chain_kwargs={"prompt": None}  # Usa template customizado
            )
            
            return True
            
        except Exception as e:
            st.error(f"Erro ao criar RAG chain: {str(e)}")
            return False
    
    def get_response(self, question: str) -> Dict[str, Any]:
        """Obtém resposta usando RAG"""
        try:
            if not self.rag_chain:
                return {
                    "answer": "❌ Sistema RAG não inicializado. Configure a API Key.",
                    "source_documents": []
                }
            
            # Faz a pergunta
            result = self.rag_chain({"question": question})
            
            # Remove raciocínio da resposta se presente
            cleaned_answer = self._remove_reasoning_from_response(result["answer"])
            
            return {
                "answer": cleaned_answer,
                "source_documents": result.get("source_documents", [])
            }
            
        except Exception as e:
            return {
                "answer": f"❌ Erro ao processar pergunta: {str(e)}",
                "source_documents": []
            }
    
    def _remove_reasoning_from_response(self, response: str) -> str:
        """Remove mostras de raciocínio interno da resposta"""
        # Padrões comuns de raciocínio a remover
        patterns_to_remove = [
            r"Vou analisar.*?(?=\n\n|\n[A-Z]|$)",
            r"Preciso entender.*?(?=\n\n|\n[A-Z]|$)",
            r"Deixe-me pensar.*?(?=\n\n|\n[A-Z]|$)",
            r"Para responder.*?(?=\n\n|\n[A-Z]|$)",
            r"Primeiro.*?(?=\n\n|\n[A-Z]|$)",
            r"Baseado no contexto.*?(?=\n\n|\n[A-Z]|$)",
            r"Analisando.*?(?=\n\n|\n[A-Z]|$)",
            r"<think>.*?</think>",
            r"<thinking>.*?</thinking>",
        ]
        
        cleaned_response = response
        for pattern in patterns_to_remove:
            cleaned_response = re.sub(pattern, "", cleaned_response, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove linhas vazias excessivas
        cleaned_response = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_response)
        
        return cleaned_response.strip()
    
    def search_relevant_content(self, query: str, k: int = 3) -> List[Document]:
        """Busca conteúdo relevante no vectorstore"""
        if not self.vectorstore:
            return []
        
        try:
            return self.vectorstore.similarity_search(query, k=k)
        except Exception as e:
            st.error(f"Erro na busca: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas do sistema"""
        return {
            "total_documents": len(self.documents),
            "vectorstore_exists": self.vectorstore is not None,
            "rag_chain_ready": self.rag_chain is not None,
            "biology_folder_exists": os.path.exists(self.biology_folder_path),
            "persist_directory": self.persist_directory
        }
    
    def clear_memory(self):
        """Limpa memória de conversação"""
        if self.memory:
            self.memory.clear()

# Instância global do sistema RAG
biology_rag_system = LocalBiologyRAG() 