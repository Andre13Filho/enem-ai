#!/usr/bin/env python3
"""
Sistema RAG Local para Língua Portuguesa
Processa documentos locais da pasta Língua Portuguesa
"""

import os
import glob
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.docstore.document import Document
from langchain.llms.base import LLM
from langchain.callbacks.manager import CallbackManagerForLLMRun

import PyPDF2
from docx import Document as DocxDocument

# Importa Groq
from groq import Groq

class GroqLLM(LLM):
    """LLM personalizado para DeepSeek R1 Distill via Groq"""
    
    api_key: str
    model_name: str = "deepseek-r1-distill-llama-70b"
    
    class Config:
        arbitrary_types_allowed = True
    
    def __init__(self, api_key: str, **kwargs):
        super().__init__(api_key=api_key, model_name="deepseek-r1-distill-llama-70b", **kwargs)
        self._client = Groq(api_key=api_key)
    
    @property
    def _llm_type(self) -> str:
        return "groq"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        try:
            response = self._client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=2048
            )
            return response.choices[0].message.content
        except Exception as e:
            from encoding_utils import safe_api_error
            return safe_api_error(e)

class LocalPortugueseRAG:
    """Sistema RAG Local para Língua Portuguesa"""
    
    def __init__(self, portuguese_folder_path: str = "./Língua Portuguesa"):
        # Configuração adaptativa para cloud/local
        try:
            from cloud_config import get_config
            self.cloud_config = get_config()
            self.portuguese_folder_path = self.cloud_config.get_documents_path("Língua Portuguesa")
            self.persist_directory = self.cloud_config.get_vectorstore_path("portuguese")
        except ImportError:
            # Fallback para configuração local
            self.cloud_config = None
            self.portuguese_folder_path = portuguese_folder_path
            self.persist_directory = "vectorstores/portuguese"
        self.persist_directory = "vectorstores/portugues"
        
        # Componentes do sistema
        self.documents: List[Document] = []
        self.embeddings = None
        self.vectorstore = None
        self.retriever = None
        self.memory = None
        self.rag_chain = None
        
        # Inicializa embeddings
        self._setup_embeddings()
    
    def _setup_embeddings(self):
        """Configura modelo de embeddings"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'}
            )
            print("✅ Embeddings inicializados")
        except Exception as e:
            print(f"❌ Erro ao inicializar embeddings: {e}")
    
    def process_portuguese_documents(self) -> bool:
        """Processa documentos da pasta Língua Portuguesa"""
        try:
            # Verifica pasta
            if not os.path.exists(self.portuguese_folder_path):
                st.error(f"❌ Pasta {self.portuguese_folder_path} não encontrada")
                return False
            
            # Lista documentos
            docx_files = glob.glob(os.path.join(self.portuguese_folder_path, "*.docx"))
            pdf_files = glob.glob(os.path.join(self.portuguese_folder_path, "*.pdf"))
            
            if not docx_files and not pdf_files:
                st.error("❌ Nenhum documento DOCX ou PDF encontrado")
                return False
            
            # Processa cada arquivo
            all_documents = []
            for file_path in docx_files + pdf_files:
                docs = self._process_single_file(file_path)
                if docs:
                    all_documents.extend(docs)
            
            if not all_documents:
                st.error("❌ Nenhum documento foi processado com sucesso")
                return False
            
            # Divide em chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", ". ", " ", ""],
                length_function=len
            )
            
            self.documents = text_splitter.split_documents(all_documents)
            
            # Cria vectorstore
            self._create_vectorstore()
            
            return True
            
        except Exception as e:
            st.error(f"❌ Erro ao processar documentos: {e}")
            return False
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Processa um único arquivo"""
        try:
            # Extrai conteúdo baseado na extensão
            if file_path.endswith('.docx'):
                content = self._extract_docx_content(file_path)
            elif file_path.endswith('.pdf'):
                content = self._extract_pdf_content(file_path)
            else:
                st.warning(f"Formato não suportado: {file_path}")
                return []
            
            if not content:
                return []
            
            # Extrai tópico do nome do arquivo
            topic = self._extract_topic_from_filename(file_path)
            
            # Cria documento
            doc = Document(
                page_content=content,
                metadata={
                    "source": file_path,
                    "topic": topic,
                    "filename": Path(file_path).name
                }
            )
            
            return [doc]
            
        except Exception as e:
            st.warning(f"Erro ao processar {file_path}: {str(e)}")
            return []
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .docx"""
        try:
            doc = DocxDocument(file_path)
            content_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            return "\n\n".join(content_parts)
        except Exception as e:
            st.warning(f"Erro ao processar DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo .pdf"""
        try:
            content_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Verifica se o PDF está criptografado
                if pdf_reader.is_encrypted:
                    # Tenta descriptografar com senha vazia
                    try:
                        pdf_reader.decrypt("")
                    except:
                        st.warning(f"PDF {Path(file_path).name} está protegido por senha - pulando...")
                        return ""
                
                # Extrai texto das páginas
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if isinstance(text, str):
                            text = text.encode('utf-8', errors='replace').decode('utf-8')
                            if text.strip():
                                content_parts.append(text)
                    except Exception as page_error:
                        st.warning(f"Erro na página {i+1} do PDF {Path(file_path).name}: {str(page_error)}")
                        continue
            
            if content_parts:
                return "\n\n".join(content_parts)
            else:
                st.info(f"PDF {Path(file_path).name} processado mas sem texto extraído")
                return ""
                
        except Exception as e:
            error_msg = str(e)
            if "PyCryptodome" in error_msg:
                st.error(f"PDF {Path(file_path).name} requer descriptografia avançada - instale: pip install PyCryptodome")
            else:
                st.warning(f"Erro ao processar PDF {Path(file_path).name}: {error_msg}")
            return ""
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extrai tópico do nome do arquivo"""
        # Remove extensão
        topic = Path(filename).stem
        
        # Limpa caracteres especiais
        topic = topic.replace("_", " ").replace("-", " ")
        
        # Capitaliza primeira letra
        return topic.title()
    
    def _create_vectorstore(self):
        """Cria ou atualiza o vectorstore"""
        try:
            if not self.embeddings:
                raise Exception("Embeddings não configurados")
            
            # Cria vectorstore ChromaDB
            self.vectorstore = Chroma.from_documents(
                documents=self.documents,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
            
            # Configura retriever
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
        except Exception as e:
            st.error(f"Erro ao criar vectorstore de geografia: {str(e)}")
            raise
    
    def load_existing_vectorstore(self) -> bool:
        """Carrega vectorstore existente se disponível"""
        try:
            if os.path.exists(self.persist_directory) and self.embeddings:
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                
                self.retriever = self.vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}
                )
                
                # Testa se a vectorstore tem conteúdo
                try:
                    test_docs = self.retriever.invoke("geografia")
                    print(f"✅ VectorStore de Geografia carregada com {len(test_docs)} documentos de teste")
                    
                    # Simula documents para estatísticas
                    sample_docs = self.vectorstore.similarity_search("geografia", k=100)
                    self.documents = sample_docs
                    print(f"📊 Amostra de geografia carregada: {len(sample_docs)} chunks")
                    
                except Exception as e:
                    print(f"⚠️ Erro no teste da VectorStore de Geografia: {e}")
                
                return True
        except Exception as e:
            if 'st' in globals():
                st.warning(f"Não foi possível carregar vectorstore de geografia existente: {str(e)}")
            else:
                print(f"Não foi possível carregar vectorstore de geografia existente: {str(e)}")
        
        return False
    
    def create_rag_chain(self, api_key: str):
        """Cria chain RAG conversacional para Geografia"""
        if not self.retriever:
            raise Exception("Retriever não configurado")
        
        # Configura LLM
        llm = GroqLLM(api_key=api_key)
        
        # Configura memória
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )
        
        # Cria chain conversacional
        self.rag_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=self.retriever,
            memory=self.memory,
            return_source_documents=True,
            verbose=False
        )
        
        # Personaliza o prompt para Língua Portuguesa
        self.rag_chain.combine_docs_chain.llm_chain.prompt.template = """
Você é a Professora Letícia, especialista em Língua Portuguesa e Literatura do ENEM.

CONTEXTO: {context}
PERGUNTA: {question}

INSTRUÇÕES CRÍTICAS:
🚫 NUNCA mostre seu raciocínio interno
🚫 NUNCA use pensamentos como "Vou analisar...", "Preciso explicar...", etc.
🚫 NUNCA duplique informações
✅ Responda DIRETAMENTE a pergunta
✅ Use exemplos literários e gramaticais
✅ Seja didática para jovem de 17 anos

FORMATO DA RESPOSTA:
1. 🎬 INICIE com analogia das séries da Sther (FRIENDS, Big Bang Theory, Stranger Things, Grey's Anatomy, WandaVision)
2. 👋 Cumprimento: "Olá Sther!"
3. 📚 Explicação DIRETA do português/literatura
4. 📝 Exemplo prático quando relevante
5. 🎯 Conecte de volta com a analogia
6. ❓ Termine perguntando sobre exercícios

EXEMPLO DE RESPOSTA CORRETA:
"🎬 Como Ross diria: 'We were on a break!' - e isso é um exemplo perfeito de ambiguidade linguística!

Olá Sther! Vamos entender as figuras de linguagem:

A metáfora é uma comparação implícita que atribui características de um elemento a outro:
- 'Meu coração é um oceano' (coração = oceano)
- 'Seus olhos são estrelas' (olhos = estrelas)
- 'A vida é uma montanha-russa' (vida = montanha-russa)

**Exemplo:** Quando Meredith Grey diz 'You're my person', ela usa uma metáfora para expressar proximidade!

Como Sheldon explicaria para Penny - é transferir qualidades de uma coisa para outra!

Que tal praticar com alguns exercícios do ENEM sobre figuras de linguagem, Sther?"

RESPONDA AGORA SEGUINDO EXATAMENTE ESTE FORMATO:"""
    
    def get_response(self, question: str) -> Dict[str, Any]:
        """Gera resposta usando RAG"""
        if not self.rag_chain:
            return {
                "answer": "Sistema RAG de Geografia não inicializado. Configure sua API key.",
                "source_documents": []
            }
        
        try:
            result = self.rag_chain({
                "question": question,
                "chat_history": []
            })
            
            # Remove raciocínio interno da resposta
            if "answer" in result:
                answer = result["answer"]
                answer = self._remove_reasoning_from_response(answer)
                
                # Aplica formatação melhorada na resposta
                from math_formatter import format_professor_response
                result["answer"] = format_professor_response(answer)
            
            return result
            
        except Exception as e:
            return {
                "answer": f"Erro ao gerar resposta: {str(e)}",
                "source_documents": []
            }
    
    def _remove_reasoning_from_response(self, response: str) -> str:
        """Remove marcadores de raciocínio interno"""
        lines = response.split('\n')
        filtered_lines = []
        skip_line = False
        
        for line in lines:
            if any(marker in line.lower() for marker in ["pensando:", "raciocínio:", "análise:", "vamos analisar"]):
                skip_line = True
                continue
            if skip_line and line.strip() == "":
                skip_line = False
                continue
            if not skip_line:
                filtered_lines.append(line)
        
        return '\n'.join(filtered_lines)
    
    def search_relevant_content(self, query: str, k: int = 3) -> List[Document]:
        """Busca conteúdo relevante"""
        if not self.retriever:
            return []
        
        try:
            docs = self.retriever.get_relevant_documents(query)
            return docs[:k]
        except Exception as e:
            st.error(f"Erro na busca: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas do sistema"""
        return {
            "total_documentos": len(self.documents) if self.documents else 0,
            "vectorstore_inicializado": self.vectorstore is not None,
            "rag_chain_configurada": self.rag_chain is not None
        }
    
    def clear_memory(self):
        """Limpa memória da conversa"""
        if self.memory:
            self.memory.clear()

# Instância global
local_portuguese_rag = LocalPortugueseRAG() 