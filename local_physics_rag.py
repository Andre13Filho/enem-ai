#!/usr/bin/env python3
"""
Sistema RAG Local para Professor Fernando - Física
Processa documentos locais da pasta física usando LangChain
"""

import streamlit as st
import os
from typing import Dict, List, Any, Optional
from pathlib import Path

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain.chains import ConversationalRetrievalChain
try:
    from langchain_community.memory import ConversationBufferMemory
except ImportError:
    from langchain.memory import ConversationBufferMemory
from langchain.llms.base import LLM
from langchain.callbacks.manager import CallbackManagerForLLMRun

# Document processors
try:
    import docx
    from pypdf import PdfReader
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False

# Groq for DeepSeek
from groq import Groq

class GroqLLM(LLM):
    """LLM personalizado para DeepSeek R1 Distill via Groq"""
    
    api_key: str
    model_name: str = "deepseek-r1-distill-llama-70b"
    
    class Config:
        arbitrary_types_allowed = True
    
    def __init__(self, api_key: str, **kwargs):
        super().__init__(api_key=api_key, model_name="deepseek-r1-distill-llama-70b", **kwargs)
        self._client = Groq(api_key=api_key)
    
    @property
    def _llm_type(self) -> str:
        return "groq"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        try:
            response = self._client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=2048
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Erro na API: {str(e)}"

class LocalPhysicsRAG:
    """Sistema RAG para documentos de física locais"""
    
    def __init__(self, physics_folder_path: str = "./física"):
        # Configuração adaptativa para cloud/local
        try:
            from cloud_config import get_config
            self.cloud_config = get_config()
            self.physics_folder_path = self.cloud_config.get_documents_path("física")
            self.persist_directory = self.cloud_config.get_vectorstore_path("physics")
        except ImportError:
            # Fallback para configuração local
            self.cloud_config = None
            self.physics_folder_path = physics_folder_path
            self.persist_directory = "./chroma_physics_vectorstore"
            
        self.vectorstore = None
        self.retriever = None
        self.memory = None
        self.rag_chain = None
        self.embeddings = None
        self.documents = []
        
        # Inicializa embeddings
        self._setup_embeddings()
    
    def _setup_embeddings(self):
        """Configura embeddings usando HuggingFace"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            print("✅ Embeddings inicializados")
        except Exception as e:
            st.error(f"Erro ao configurar embeddings: {str(e)}")
    
    def process_physics_documents(self) -> bool:
        """Processa todos os documentos da pasta física"""
        if not os.path.exists(self.physics_folder_path):
            st.error(f"Pasta não encontrada: {self.physics_folder_path}")
            return False
        
        if not DOCUMENT_PROCESSING_AVAILABLE:
            st.error("Bibliotecas de processamento não disponíveis. Instale: pip install python-docx pypdf")
            return False
        
        try:
            # Lista todos os arquivos na pasta física
            physics_files = []
            for root, dirs, files in os.walk(self.physics_folder_path):
                for file in files:
                    if file.lower().endswith(('.docx', '.pdf', '.txt')):
                        physics_files.append(os.path.join(root, file))
            
            st.info(f"Encontrados {len(physics_files)} arquivos para processar...")
            
            # Processa cada arquivo
            all_documents = []
            progress_bar = st.progress(0)
            
            for i, file_path in enumerate(physics_files):
                try:
                    docs = self._process_single_file(file_path)
                    all_documents.extend(docs)
                    progress_bar.progress((i + 1) / len(physics_files))
                except Exception as e:
                    st.warning(f"Erro ao processar {file_path}: {str(e)}")
                    continue
            
            self.documents = all_documents
            
            if not self.documents:
                st.error("Nenhum documento foi processado com sucesso")
                return False
            
            # Cria vectorstore
            self._create_vectorstore()
            
            st.success(f"✅ Processados {len(self.documents)} chunks de {len(physics_files)} arquivos")
            return True
            
        except Exception as e:
            st.error(f"Erro no processamento: {str(e)}")
            return False
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Processa um único arquivo"""
        file_extension = Path(file_path).suffix.lower()
        file_name = Path(file_path).name
        
        # Extrai conteúdo baseado no tipo
        if file_extension == '.docx':
            content = self._extract_docx_content(file_path)
        elif file_extension == '.pdf':
            content = self._extract_pdf_content(file_path)
        elif file_extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
        else:
            return []
        
        if not content.strip():
            return []
        
        # Cria documento LangChain
        doc = Document(
            page_content=content,
            metadata={
                "source": file_name,
                "file_path": file_path,
                "topic": self._extract_topic_from_filename(file_name),
                "subject": "Física"
            }
        )
        
        # Divide em chunks menores
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = text_splitter.split_documents([doc])
        return chunks
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo DOCX"""
        try:
            doc = docx.Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            return "\n".join(content)
        except Exception as e:
            st.warning(f"Erro ao processar DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivo PDF"""
        try:
            reader = PdfReader(file_path)
            content = []
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    content.append(text.strip())
            return "\n".join(content)
        except Exception as e:
            st.warning(f"Erro ao processar PDF {file_path}: {str(e)}")
            return ""
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extrai tópico do nome do arquivo"""
        topic = Path(filename).stem
        topic = topic.replace("_", " ").replace("-", " ")
        return topic.title()
    
    def _create_vectorstore(self):
        """Cria vectorstore com fallback para ambiente cloud"""
        if not self.documents:
            raise ValueError("Nenhum documento disponível para criar vectorstore")
        
        try:
            # Primeira tentativa: vectorstore persistente
            try:
                os.makedirs(self.persist_directory, exist_ok=True)
                self.vectorstore = Chroma.from_documents(
                    documents=self.documents,
                    embedding=self.embeddings,
                    persist_directory=self.persist_directory
                )
                self.vectorstore.persist()
                print(f"✅ Vectorstore de física persistente criado em {self.persist_directory}")
            except Exception as persist_error:
                print(f"⚠️ Falha na criação persistente de física: {persist_error}")
                # Segunda tentativa: vectorstore em memória
                self.vectorstore = Chroma.from_documents(
                    documents=self.documents,
                    embedding=self.embeddings
                )
                print("✅ Vectorstore de física em memória criado com sucesso")
            
            # Configura retriever
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
            # Teste básico de funcionamento
            test_docs = self.retriever.invoke("física")
            print(f"✅ Vectorstore de física funcionando - teste retornou {len(test_docs)} documentos")
            
        except Exception as e:
            print(f"❌ Erro crítico na criação do vectorstore de física: {str(e)}")
            raise
    
    def load_existing_vectorstore(self) -> bool:
        """Carrega vectorstore existente ou cria um novo em memória"""
        try:
            # Primeiro tenta carregar vectorstore persistente
            if os.path.exists(self.persist_directory):
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                self.retriever = self.vectorstore.as_retriever(search_kwargs={"k": 5})
                print("✅ Vectorstore persistente de física carregado com sucesso")
                return True
        except Exception as e:
            print(f"⚠️ Erro ao carregar vectorstore persistente de física: {str(e)}")
        
        # Se falhar, tenta criar vectorstore em memória
        try:
            # Processa documentos se ainda não foram processados
            if not self.documents:
                print("🔄 Processando documentos de física para vectorstore em memória...")
                if not self.process_physics_documents():
                    return False
            
            if self.documents:
                # Cria vectorstore em memória (sem persistência)
                self.vectorstore = Chroma.from_documents(
                    documents=self.documents,
                    embedding=self.embeddings
                )
                self.retriever = self.vectorstore.as_retriever(search_kwargs={"k": 5})
                print("✅ Vectorstore de física em memória criado com sucesso")
                return True
            else:
                print("❌ Nenhum documento de física disponível para criar vectorstore")
                return False
                
        except Exception as e:
            print(f"❌ Erro ao criar vectorstore de física em memória: {str(e)}")
            return False
    
    def create_rag_chain(self, api_key: str):
        """Cria chain RAG conversacional"""
        try:
            if not self.retriever:
                st.error("Retriever não configurado")
                return
            
            # Configura LLM
            llm = GroqLLM(api_key=api_key)
            
            # Configura memória conversacional
            self.memory = ConversationBufferMemory(
                memory_key="chat_history",
                return_messages=True,
                output_key="answer"
            )
            
            # Cria chain conversacional
            self.rag_chain = ConversationalRetrievalChain.from_llm(
                llm=llm,
                retriever=self.retriever,
                memory=self.memory,
                return_source_documents=True,
                verbose=False
            )
            
            print("✅ RAG Chain configurado")
            
        except Exception as e:
            st.error(f"Erro ao criar RAG chain: {str(e)}")
    
    def get_response(self, question: str) -> Dict[str, Any]:
        """Gera resposta usando RAG"""
        try:
            if not self.rag_chain:
                return {
                    "answer": "Sistema RAG não inicializado",
                    "source_documents": []
                }
            
            # Adiciona contexto específico de física
            physics_context = """
Você é o Professor Fernando, especialista em Física para o ENEM. 

INSTRUÇÕES CRÍTICAS:
🚫 NUNCA mostre seu raciocínio interno
🚫 NUNCA use pensamentos como "Vou analisar...", "Preciso calcular...", etc.
🚫 NUNCA duplique informações
✅ Responda DIRETAMENTE a pergunta
✅ Use fórmulas em LaTeX: $ ou $$
✅ Seja didático para jovem de 17 anos

FORMATO DA RESPOSTA:
1. 🎬 INICIE com analogia das séries da Sther (FRIENDS, Big Bang Theory, Stranger Things, Grey's Anatomy, WandaVision)
2. 👋 Cumprimento: "Olá Sther!"
3. 📚 Explicação DIRETA da física
4. 📝 Exemplo prático quando relevante
5. 🎯 Conecte de volta com a analogia
6. ❓ Termine perguntando sobre exercícios

EXEMPLO DE FÓRMULAS:
- Velocidade: $v = \\frac{\\Delta s}{\\Delta t}$
- Segunda Lei de Newton: $$F = ma$$
- Energia cinética: $$E_c = \\frac{mv^2}{2}$$

Responda AGORA seguindo EXATAMENTE este formato.
"""
            
            enhanced_question = f"{physics_context}\n\nPergunta: {question}"
            
            result = self.rag_chain({"question": enhanced_question})
            
            # Aplica formatação melhorada na resposta
            if "answer" in result:
                try:
                    from physics_formatter import format_professor_response
                    result["answer"] = format_professor_response(result["answer"])
                except ImportError:
                    # Se o formatador não estiver disponível, usa formatação básica
                    answer = result["answer"]
                    # Remove raciocínio básico
                    import re
                    reasoning_patterns = [
                        r'Vou calcular.*?(?=\n|$)',
                        r'Primeiro.*?vamos.*?(?=\n|$)',
                        r'Pensando.*?(?=\n|$)',
                        r'Analisando.*?(?=\n|$)',
                    ]
                    for pattern in reasoning_patterns:
                        answer = re.sub(pattern, '', answer, flags=re.IGNORECASE)
                    result["answer"] = answer.strip()
            
            return {
                "answer": result.get("answer", "Não foi possível gerar resposta"),
                "source_documents": result.get("source_documents", [])
            }
            
        except Exception as e:
            return {
                "answer": f"Erro ao processar pergunta: {str(e)}",
                "source_documents": []
            }
    
    def search_relevant_content(self, query: str, k: int = 3) -> List[Document]:
        """Busca conteúdo relevante"""
        try:
            if not self.retriever:
                return []
            
            docs = self.retriever.get_relevant_documents(query)
            return docs[:k]
        except Exception as e:
            print(f"Erro na busca: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas do sistema"""
        stats = {
            "documents_processed": len(self.documents),
            "vectorstore_exists": self.vectorstore is not None,
            "rag_chain_ready": self.rag_chain is not None,
            "physics_folder": self.physics_folder_path,
            "persist_directory": self.persist_directory
        }
        
        return stats
    
    def clear_memory(self):
        """Limpa memória conversacional"""
        if self.memory:
            self.memory.clear()

# Instância global
local_physics_rag = LocalPhysicsRAG() 