# -*- coding: utf-8 -*-
"""
Sistema de Análise de Redação - Professora Carla
Análise completa baseada nos critérios do ENEM
"""

import streamlit as st
import os
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict
import tempfile

# Importações para processamento de PDF
try:
    import PyPDF2
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ProfessorRedacao:
    def __init__(self, base_path: str = "redação/", success_cases_path: str = "cases_sucesso_redacao/"):
        self.base_path = Path(base_path)
        self.success_cases_path = Path(success_cases_path)
        self.criterios_enem = {}
        self.redacoes_nota_1000 = []
        
        # Carrega os dados na inicialização
        self._load_criterios_enem()
        self._load_redacoes_nota_1000()

    def _load_criterios_enem(self):
        """Carrega os critérios de avaliação do ENEM dos arquivos .docx"""
        if not DOCX_AVAILABLE:
            return
            
        criterios_files = {
            "coesao": "Coesão.docx",
            "coerencia": "Coerência.docx", 
            "argumentacao": "Argumentação.docx",
            "gramatica": "Gramática.docx",
            "proposta_intervencao": "Elementos da Proposta de intervenção.docx"
        }
        
        for criterio, filename in criterios_files.items():
            file_path = self.base_path / filename
            if file_path.exists():
                try:
                    doc = Document(file_path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    self.criterios_enem[criterio] = content
                except Exception as e:
                    print(f"Erro ao carregar {filename}: {e}")

    def _load_redacoes_nota_1000(self):
        """Carrega exemplos de redações nota 1000"""
        if not DOCX_AVAILABLE:
            return
            
        for file_path in self.success_cases_path.glob("*.docx"):
            try:
                doc = Document(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                self.redacoes_nota_1000.append({
                    "nome": file_path.stem,
                    "conteudo": content
                })
        except Exception as e:
                print(f"Erro ao carregar {file_path.name}: {e}")

    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Extrai texto do PDF usando múltiplas estratégias"""
        text = ""
        
        # Estratégia 1: PyPDF2 para PDFs com texto
        if PDF_AVAILABLE:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                    tmp_file.write(pdf_content)
                    tmp_file.flush()
                    
                    with open(tmp_file.name, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                                text += page_text + "\n"
                    
                    os.unlink(tmp_file.name)
                    
                if text.strip():
                    return text
            except Exception:
                pass

        # Estratégia 2: PyMuPDF para PDFs mais complexos
        if PDF_AVAILABLE:
            try:
                pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                page_text = page.get_text()
                if page_text.strip():
                        text += page_text + "\n"
                pdf_document.close()
                
                if text.strip():
                    return text
            except Exception:
                pass

        return text if text.strip() else "Não foi possível extrair texto do PDF."

    def analyze_redacao_pdf(self, pdf_content: bytes, filename: str) -> str:
        """Analisa a redação do PDF e retorna feedback detalhado"""
        
        # Extrai o texto do PDF
        texto_redacao = self.extract_text_from_pdf(pdf_content)
        
        if not texto_redacao or texto_redacao == "Não foi possível extrair texto do PDF.":
            return self._generate_error_analysis()
        
        # Realiza a análise completa
        analise = self._analyze_text(texto_redacao, filename)
        
        return analise

    def _analyze_text(self, texto: str, filename: str) -> str:
        """Realiza análise detalhada do texto da redação"""
        
        # Análises básicas
        palavras = len(texto.split())
        paragrafos = len([p for p in texto.split('\n\n') if p.strip()])
        linhas = len([l for l in texto.split('\n') if l.strip()])
        
        # Análise estrutural
        estrutura_score = self._analyze_structure(texto)
        
        # Análise de conteúdo
        conteudo_score = self._analyze_content(texto)
        
        # Análise de linguagem
        linguagem_score = self._analyze_language(texto)
        
        # Análise de proposta de intervenção
        intervencao_score = self._analyze_intervention(texto)
        
        # Cálculo da nota final
        nota_final = self._calculate_final_score(estrutura_score, conteudo_score, linguagem_score, intervencao_score)
        
        # Gera o relatório completo
        relatorio = self._generate_detailed_report(
            filename, texto, palavras, paragrafos, linhas,
            estrutura_score, conteudo_score, linguagem_score, intervencao_score, nota_final
        )
        
        return relatorio

    def _analyze_structure(self, texto: str) -> Dict:
        """Analisa a estrutura da redação"""
        paragrafos = [p.strip() for p in texto.split('\n\n') if p.strip()]
        
        score = {
            "introducao": 0,
            "desenvolvimento": 0,
            "conclusao": 0,
            "total": 0,
            "feedback": []
        }
        
        if len(paragrafos) >= 4:
            score["introducao"] = 50
            score["desenvolvimento"] = 40
            score["conclusao"] = 30
            score["feedback"].append("✅ Estrutura adequada com introdução, desenvolvimento e conclusão")
        elif len(paragrafos) == 3:
            score["introducao"] = 40
            score["desenvolvimento"] = 30
            score["conclusao"] = 20
            score["feedback"].append("⚠️ Estrutura básica presente, mas poderia ter mais parágrafos de desenvolvimento")
        else:
            score["introducao"] = 20
            score["desenvolvimento"] = 10
            score["conclusao"] = 10
            score["feedback"].append("❌ Estrutura inadequada - redação muito curta")
        
        # Verifica elementos específicos
        if any(palavra in texto.lower() for palavra in ["portanto", "assim", "dessa forma", "logo"]):
            score["conclusao"] += 20
            score["feedback"].append("✅ Conectivos de conclusão identificados")
        
        score["total"] = score["introducao"] + score["desenvolvimento"] + score["conclusao"]
        return score

    def _analyze_content(self, texto: str) -> Dict:
        """Analisa o conteúdo e argumentação"""
        score = {
            "argumentacao": 0,
            "repertorio": 0,
            "coerencia": 0,
            "total": 0,
            "feedback": []
        }
        
        # Verifica argumentação
        argumentos = len(re.findall(r'\b(porque|pois|uma vez que|visto que|já que|dado que)\b', texto.lower()))
        if argumentos >= 3:
            score["argumentacao"] = 60
            score["feedback"].append("✅ Boa quantidade de conectivos argumentativos")
        elif argumentos >= 1:
            score["argumentacao"] = 40
            score["feedback"].append("⚠️ Argumentação presente, mas pode ser melhorada")
        else:
            score["argumentacao"] = 20
            score["feedback"].append("❌ Argumentação insuficiente")
        
        # Verifica repertório sociocultural
        indicadores_repertorio = ["segundo", "de acordo com", "conforme", "dados", "pesquisa", "estudo"]
        repertorio_count = sum(1 for ind in indicadores_repertorio if ind in texto.lower())
        
        if repertorio_count >= 2:
            score["repertorio"] = 40
            score["feedback"].append("✅ Bom uso de repertório sociocultural")
        elif repertorio_count >= 1:
            score["repertorio"] = 25
            score["feedback"].append("⚠️ Repertório presente, mas pode ser ampliado")
        else:
            score["repertorio"] = 10
            score["feedback"].append("❌ Repertório sociocultural insuficiente")
        
        # Coerência básica
        score["coerencia"] = 40 if len(texto.split()) > 150 else 20
        
        score["total"] = score["argumentacao"] + score["repertorio"] + score["coerencia"]
        return score

    def _analyze_language(self, texto: str) -> Dict:
        """Analisa aspectos linguísticos"""
        score = {
            "coesao": 0,
            "registro": 0,
            "variedade": 0,
            "total": 0,
            "feedback": []
        }
        
        # Conectivos de coesão
        conectivos = len(re.findall(r'\b(além disso|por outro lado|entretanto|contudo|todavia|portanto|assim|dessa forma)\b', texto.lower()))
        if conectivos >= 3:
            score["coesao"] = 40
            score["feedback"].append("✅ Bom uso de conectivos coesivos")
        elif conectivos >= 1:
            score["coesao"] = 25
            score["feedback"].append("⚠️ Uso básico de conectivos")
        else:
            score["coesao"] = 10
            score["feedback"].append("❌ Poucos conectivos coesivos")
        
        # Registro formal
        if not re.search(r'\b(né|tipo|meio que|daí|aí)\b', texto.lower()):
            score["registro"] = 30
            score["feedback"].append("✅ Registro formal adequado")
        else:
            score["registro"] = 15
            score["feedback"].append("⚠️ Presença de marcas de oralidade")
        
        # Variedade lexical
        palavras_unicas = len(set(texto.lower().split()))
        total_palavras = len(texto.split())
        if total_palavras > 0 and (palavras_unicas / total_palavras) > 0.6:
            score["variedade"] = 30
            score["feedback"].append("✅ Boa variedade lexical")
        else:
            score["variedade"] = 15
            score["feedback"].append("⚠️ Variedade lexical pode ser melhorada")
        
        score["total"] = score["coesao"] + score["registro"] + score["variedade"]
        return score

    def _analyze_intervention(self, texto: str) -> Dict:
        """Analisa a proposta de intervenção"""
        score = {
            "agente": 0,
            "acao": 0,
            "meio": 0,
            "finalidade": 0,
            "detalhamento": 0,
            "total": 0,
            "feedback": []
        }
        
        # Verifica presença de agente
        agentes = ["governo", "estado", "ministério", "secretaria", "escola", "família", "sociedade", "mídia"]
        if any(agente in texto.lower() for agente in agentes):
            score["agente"] = 40
            score["feedback"].append("✅ Agente da intervenção identificado")
        else:
            score["agente"] = 10
            score["feedback"].append("❌ Agente da intervenção não identificado")
        
        # Verifica ação
        acoes = ["deve", "precisa", "necessário", "implementar", "criar", "desenvolver", "promover"]
        if any(acao in texto.lower() for acao in acoes):
            score["acao"] = 40
            score["feedback"].append("✅ Ação proposta identificada")
        else:
            score["acao"] = 10
            score["feedback"].append("❌ Ação não claramente proposta")
        
        # Verifica meio/modo
        meios = ["através", "por meio", "mediante", "via", "utilizando", "campanhas", "programas"]
        if any(meio in texto.lower() for meio in meios):
            score["meio"] = 20
            score["feedback"].append("✅ Meio/modo da intervenção presente")
        else:
            score["meio"] = 5
            score["feedback"].append("⚠️ Meio/modo da intervenção pouco detalhado")
        
        score["total"] = score["agente"] + score["acao"] + score["meio"] + score["finalidade"] + score["detalhamento"]
        return score

    def _calculate_final_score(self, estrutura: Dict, conteudo: Dict, linguagem: Dict, intervencao: Dict) -> int:
        """Calcula a nota final baseada nos critérios do ENEM"""
        
        # Competência 1: Estrutura (0-200)
        comp1 = min(estrutura["total"] * 2, 200)
        
        # Competência 2: Conteúdo (0-200)
        comp2 = min(conteudo["total"] * 1.5, 200)
        
        # Competência 3: Linguagem (0-200)
        comp3 = min(linguagem["total"] * 2, 200)
        
        # Competência 4: Argumentação (já incluída no conteúdo)
        comp4 = min(conteudo["argumentacao"] * 3, 200)
        
        # Competência 5: Intervenção (0-200)
        comp5 = min(intervencao["total"] * 2, 200)
        
        # Nota final (média das competências)
        nota_final = (comp1 + comp2 + comp3 + comp4 + comp5) / 5
        
        return int(nota_final)

    def _generate_detailed_report(self, filename: str, texto: str, palavras: int, paragrafos: int, linhas: int,
                                estrutura: Dict, conteudo: Dict, linguagem: Dict, intervencao: Dict, nota_final: int) -> str:
        """Gera relatório detalhado da análise"""
        
        relatorio = f"""
# 📝 **Análise Detalhada da Redação**
**Arquivo:** {filename}
**Data da Análise:** {datetime.now().strftime("%d/%m/%Y às %H:%M")}

---

## 📊 **Estatísticas Gerais**
- **Palavras:** {palavras}
- **Parágrafos:** {paragrafos}
- **Linhas:** {linhas}

---

## 🎯 **NOTA FINAL: {nota_final}/1000**

---

## 📋 **Análise por Competências**

### 🏗️ **Competência 1 - Estrutura Textual**
**Pontuação:** {estrutura['total']}/120

**Feedback:**
"""
        
        for feedback in estrutura['feedback']:
            relatorio += f"\n- {feedback}"
        
        relatorio += f"""

### 💭 **Competência 2 - Conteúdo e Argumentação**
**Pontuação:** {conteudo['total']}/140

**Feedback:**
"""
        
        for feedback in conteudo['feedback']:
            relatorio += f"\n- {feedback}"
        
        relatorio += f"""

### 🗣️ **Competência 3 - Aspectos Linguísticos**
**Pontuação:** {linguagem['total']}/100

**Feedback:**
"""
        
        for feedback in linguagem['feedback']:
            relatorio += f"\n- {feedback}"
        
        relatorio += f"""

### 🎯 **Competência 5 - Proposta de Intervenção**
**Pontuação:** {intervencao['total']}/100

**Feedback:**
"""
        
        for feedback in intervencao['feedback']:
            relatorio += f"\n- {feedback}"
        
        relatorio += """

---

## 💡 **Sugestões de Melhoria**

### 📚 **Para a próxima redação:**
1. **Estrutura:** Mantenha sempre 4-5 parágrafos bem definidos
2. **Argumentação:** Use mais dados, estatísticas e referências
3. **Coesão:** Varie os conectivos para melhor fluidez
4. **Intervenção:** Detalhe mais os elementos da proposta

### 🎓 **Dicas da Professora Carla:**
- Leia redações nota 1000 para se inspirar
- Pratique a escrita de propostas de intervenção detalhadas
- Amplie seu repertório sociocultural com leituras diversas
- Revise sempre a gramática e a coesão textual

---

**✨ Continue praticando! Cada redação é um passo mais próximo da nota 1000! ✨**
"""
        
        return relatorio

    def _generate_error_analysis(self) -> str:
        """Gera análise para casos de erro na extração"""
        return """
# ❌ **Erro na Análise da Redação**

Não foi possível extrair o texto do PDF enviado.

## 🔧 **Possíveis Soluções:**
1. Verifique se o arquivo é um PDF válido
2. Certifique-se de que o PDF não está protegido por senha
3. Tente converter o arquivo para um formato mais simples
4. Se o PDF contém apenas imagens, use um OCR antes de enviar

## 📞 **Precisa de Ajuda?**
Entre em contato com a Professora Carla para assistência técnica.
"""

    def get_success_cases(self) -> List[str]:
        """Retorna lista de casos de sucesso disponíveis"""
        return [redacao["nome"] for redacao in self.redacoes_nota_1000]

# Instância global
professor_redacao = ProfessorRedacao()

def analyze_redacao_pdf(pdf_content: bytes, filename: str) -> str:
    """Função wrapper para análise de redação"""
    return professor_redacao.analyze_redacao_pdf(pdf_content, filename)

def setup_redacao_ui():
    """Configura a interface do sistema de redação"""
    st.markdown("""
    <div class="teacher-intro">
        <h3>✍️ Professora Carla - Análise de Redação</h3>
        <p>Sistema completo de análise baseado nos critérios do ENEM</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### 📤 **Envie sua Redação**")
    
    uploaded_file = st.file_uploader(
        "Escolha um arquivo PDF com sua redação:",
        type=['pdf'],
        help="Envie sua redação em formato PDF para análise completa"
    )
    
    if uploaded_file is not None:
        if st.button("🔍 Analisar Redação", type="primary"):
            with st.spinner("📝 Professora Carla analisando sua redação..."):
                try:
                    # Lê o conteúdo do arquivo
                    pdf_content = uploaded_file.read()
                    
                    # Analisa a redação
                    analise = analyze_redacao_pdf(pdf_content, uploaded_file.name)
                    
                    # Exibe o resultado
                    st.markdown("### 📋 **Resultado da Análise**")
                    st.markdown(analise)
                    
                    # Botão para download do relatório
                    st.download_button(
                        label="📥 Baixar Relatório Completo",
                        data=analise,
                        file_name=f"analise_redacao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                        mime="text/markdown"
                    )
                    
                except Exception as e:
                    st.error(f"❌ Erro ao processar a redação: {str(e)}")
                    st.info("💡 Verifique se o arquivo é um PDF válido e tente novamente.")
    
    # Informações adicionais
    with st.expander("ℹ️ Como funciona a análise?"):
        st.markdown("""
        **A Professora Carla analisa sua redação baseada nos 5 critérios do ENEM:**
        
        1. **🏗️ Estrutura Textual** - Organização e formato dissertativo-argumentativo
        2. **💭 Conteúdo** - Argumentação e repertório sociocultural  
        3. **🗣️ Linguagem** - Coesão, registro formal e variedade lexical
        4. **🎯 Argumentação** - Desenvolvimento lógico das ideias
        5. **📋 Proposta de Intervenção** - Detalhamento e viabilidade
        
        **📊 Você receberá:**
        - Nota de 0 a 1000 pontos
        - Feedback detalhado por competência
        - Sugestões específicas de melhoria
        - Dicas personalizadas da Professora Carla
        """)
    
    # Casos de sucesso
    if professor_redacao.redacoes_nota_1000:
        with st.expander("🏆 Exemplos de Redações Nota 1000"):
            st.markdown("**Inspire-se com estes exemplos:**")
            for redacao in professor_redacao.redacoes_nota_1000[:3]:  # Mostra apenas os 3 primeiros
                st.markdown(f"- 📝 {redacao['nome']}") 