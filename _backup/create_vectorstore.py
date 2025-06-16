#!/usr/bin/env python3
"""
Script para criar VectorStore com ChromaDB e Embeddings HuggingFace
Processa todos os documentos da pasta matemática para criar sistema RAG
"""

import os
import sys
from pathlib import Path
import time
from typing import List, Dict, Any

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain.schema import Document

# Embeddings
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

# Processadores de documentos
try:
    import docx
    from pypdf import PdfReader
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    print("❌ Instale as dependências: pip install python-docx pypdf")
    sys.exit(1)

class MathVectorStoreCreator:
    """Criador de VectorStore para documentos de matemática"""
    
    def __init__(self, math_folder: str = "./matemática", persist_dir: str = "./chroma_math_vectorstore"):
        self.math_folder = Path(math_folder)
        self.persist_dir = persist_dir
        self.embeddings = None
        self.vectorstore = None
        self.documents = []
        
        # Configuração do text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        print(f"🎯 Configuração:")
        print(f"   📁 Pasta origem: {self.math_folder}")
        print(f"   💾 Diretório persistência: {self.persist_dir}")
    
    def setup_embeddings(self):
        """Configura embeddings HuggingFace"""
        print("\n🧠 Configurando embeddings HuggingFace...")
        
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            print("✅ Embeddings configurados: sentence-transformers/all-MiniLM-L6-v2")
            
            # Teste rápido
            test_embedding = self.embeddings.embed_query("teste")
            print(f"✅ Teste de embedding: dimensão {len(test_embedding)}")
            
        except Exception as e:
            print(f"❌ Erro ao configurar embeddings: {e}")
            sys.exit(1)
    
    def find_math_documents(self) -> List[Path]:
        """Encontra todos os documentos de matemática"""
        print(f"\n📂 Buscando documentos em {self.math_folder}...")
        
        if not self.math_folder.exists():
            print(f"❌ Pasta não encontrada: {self.math_folder}")
            return []
        
        # Busca arquivos suportados
        supported_extensions = ['.docx', '.pdf', '.txt']
        documents = []
        
        for ext in supported_extensions:
            files = list(self.math_folder.glob(f"*{ext}"))
            documents.extend(files)
            print(f"   📄 {ext.upper()}: {len(files)} arquivos")
        
        print(f"📊 Total: {len(documents)} documentos encontrados")
        return documents
    
    def extract_content_from_docx(self, file_path: Path) -> str:
        """Extrai conteúdo de arquivo DOCX"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            print(f"⚠️  Erro ao processar {file_path.name}: {e}")
            return ""
    
    def extract_content_from_pdf(self, file_path: Path) -> str:
        """Extrai conteúdo de arquivo PDF"""
        try:
            reader = PdfReader(file_path)
            content_parts = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text.strip())
                except Exception as e:
                    print(f"⚠️  Erro na página {page_num + 1} de {file_path.name}: {e}")
                    continue
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            print(f"⚠️  Erro ao processar {file_path.name}: {e}")
            return ""
    
    def process_single_document(self, file_path: Path) -> List[Document]:
        """Processa um único documento"""
        file_extension = file_path.suffix.lower()
        file_name = file_path.name
        
        print(f"   📄 Processando: {file_name}")
        
        # Extrai conteúdo baseado no tipo
        if file_extension == '.docx':
            content = self.extract_content_from_docx(file_path)
        elif file_extension == '.pdf':
            content = self.extract_content_from_pdf(file_path)
        else:
            print(f"⚠️  Tipo não suportado: {file_extension}")
            return []
        
        if not content.strip():
            print(f"⚠️  Conteúdo vazio: {file_name}")
            return []
        
        # Identifica tópico do nome do arquivo
        topic = file_name.replace('.docx', '').replace('.pdf', '').replace('_', ' ')
        
        # Cria documento LangChain
        doc = Document(
            page_content=content,
            metadata={
                "source": file_name,
                "file_path": str(file_path),
                "type": "matematica",
                "topic": topic,
                "extension": file_extension
            }
        )
        
        # Divide em chunks
        chunks = self.text_splitter.split_documents([doc])
        
        print(f"      ✅ {len(chunks)} chunks criados")
        return chunks
    
    def process_all_documents(self):
        """Processa todos os documentos"""
        print("\n🔄 Processando todos os documentos...")
        
        # Encontra documentos
        document_files = self.find_math_documents()
        
        if not document_files:
            print("❌ Nenhum documento encontrado!")
            return False
        
        # Processa cada documento
        all_chunks = []
        successful_files = 0
        
        print(f"\n📚 Processando {len(document_files)} documentos...")
        
        for i, file_path in enumerate(document_files):
            print(f"[{i+1}/{len(document_files)}] Processando {file_path.name}")
            try:
                chunks = self.process_single_document(file_path)
                if chunks:
                    all_chunks.extend(chunks)
                    successful_files += 1
            except Exception as e:
                print(f"❌ Erro ao processar {file_path.name}: {e}")
                continue
        
        self.documents = all_chunks
        
        print(f"\n📊 Resultado do processamento:")
        print(f"   📄 Arquivos processados: {successful_files}/{len(document_files)}")
        print(f"   📝 Total de chunks: {len(all_chunks)}")
        
        if not all_chunks:
            print("❌ Nenhum chunk foi criado!")
            return False
        
        return True
    
    def create_vectorstore(self):
        """Cria VectorStore com ChromaDB"""
        print("\n💾 Criando VectorStore com ChromaDB...")
        
        if not self.documents:
            print("❌ Nenhum documento processado!")
            return False
        
        if not self.embeddings:
            print("❌ Embeddings não configurados!")
            return False
        
        try:
            # Remove diretório existente se houver
            if os.path.exists(self.persist_dir):
                import shutil
                shutil.rmtree(self.persist_dir)
                print(f"🗑️  Removido diretório existente: {self.persist_dir}")
            
            # Cria VectorStore
            print("🔄 Criando embeddings e armazenando no ChromaDB...")
            
            self.vectorstore = Chroma.from_documents(
                documents=self.documents,
                embedding=self.embeddings,
                persist_directory=self.persist_dir
            )
            
            print(f"✅ VectorStore criado em: {self.persist_dir}")
            print(f"📊 Documentos armazenados: {len(self.documents)}")
            
            return True
            
        except Exception as e:
            print(f"❌ Erro ao criar VectorStore: {e}")
            return False
    
    def test_vectorstore(self):
        """Testa o VectorStore criado"""
        print("\n🧪 Testando VectorStore...")
        
        if not self.vectorstore:
            print("❌ VectorStore não foi criado!")
            return False
        
        # Testes de busca
        test_queries = [
            "função quadrática",
            "trigonometria",
            "geometria",
            "logaritmo",
            "progressão aritmética"
        ]
        
        retriever = self.vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 3}
        )
        
        for query in test_queries:
            print(f"\n🔍 Testando busca: '{query}'")
            try:
                docs = retriever.get_relevant_documents(query)
                print(f"   📊 Encontrados: {len(docs)} documentos relevantes")
                
                for i, doc in enumerate(docs[:2]):  # Mostra apenas os 2 primeiros
                    source = doc.metadata.get('source', 'Desconhecido')
                    topic = doc.metadata.get('topic', 'Geral')
                    content_preview = doc.page_content[:100] + "..." if len(doc.page_content) > 100 else doc.page_content
                    
                    print(f"      {i+1}. {topic} ({source})")
                    print(f"         {content_preview}")
                
            except Exception as e:
                print(f"   ❌ Erro na busca: {e}")
        
        return True

def main():
    """Função principal"""
    print("🎓 ENEM AI Helper - Criador de VectorStore com ChromaDB")
    print("=" * 60)
    
    # Verifica se a pasta matemática existe
    math_folder = Path("./matemática")
    if not math_folder.exists():
        print(f"❌ Pasta não encontrada: {math_folder}")
        print("💡 Certifique-se de que a pasta 'matemática' existe com documentos")
        return
    
    # Cria instância do criador
    creator = MathVectorStoreCreator()
    
    try:
        # 1. Configura embeddings
        creator.setup_embeddings()
        
        # 2. Processa documentos
        success = creator.process_all_documents()
        if not success:
            print("❌ Falha no processamento dos documentos")
            return
        
        # 3. Cria VectorStore
        success = creator.create_vectorstore()
        if not success:
            print("❌ Falha na criação da VectorStore")
            return
        
        # 4. Testa VectorStore
        creator.test_vectorstore()
        
        print("\n🎉 VectorStore criada com sucesso!")
        print("\n🚀 Para usar no sistema principal:")
        print("1. Execute: streamlit run app.py")
        print("2. Configure sua API Key do OpenRouter")
        print("3. Selecione 'Matemática'")
        print("4. O sistema carregará automaticamente a VectorStore criada!")
        
    except KeyboardInterrupt:
        print("\n⏹️  Processo interrompido pelo usuário")
    except Exception as e:
        print(f"\n❌ Erro inesperado: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 