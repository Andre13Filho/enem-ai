"""
Sistema RAG Local para Professor Carlos - Matemática
Processa documentos locais da pasta matemática usando LangChain
"""

import streamlit as st
import os
from typing import Dict, List, Any, Optional
from pathlib import Path
import tempfile
import hashlib

# LangChain imports
from langchain_community.document_loaders import TextLoader, DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain.chains import ConversationalRetrievalChain
try:
    from langchain_community.memory import ConversationBufferMemory
except ImportError:
    from langchain.memory import ConversationBufferMemory
from langchain.llms.base import LLM
from langchain.schema import BaseMessage, AIMessage, HumanMessage
from langchain.callbacks.manager import CallbackManagerForLLMRun

# Document processors
try:
    import docx
    from pypdf import PdfReader
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False

# OpenAI for DeepSeek
from groq import Groq

class GroqLLM(LLM):
    """LLM personalizado para DeepSeek R1 Distill via Groq"""
    
    api_key: str
    model_name: str = "deepseek-r1-distill-llama-70b"
    
    class Config:
        arbitrary_types_allowed = True
    
    def __init__(self, api_key: str, **kwargs):
        super().__init__(api_key=api_key, model_name="deepseek-r1-distill-llama-70b", **kwargs)
        self._client = Groq(api_key=api_key)
    
    @property
    def _llm_type(self) -> str:
        return "groq"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        try:
            response = self._client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=2048
            )
            return response.choices[0].message.content
        except Exception as e:
            from encoding_utils import safe_api_error
            return safe_api_error(e)

class LocalMathRAG:
    """Sistema RAG para documentos de matemática locais"""
    
    def __init__(self, math_folder_path: str = "./matemática"):
        self.math_folder_path = math_folder_path
        self.vectorstore = None
        self.retriever = None
        self.memory = None
        self.rag_chain = None
        self.embeddings = None
        self.documents = []
        self.processed_files = set()
        
        # Configuração do ChromaDB local
        self.persist_directory = "./chroma_math_vectorstore"
        
        # Inicializa embeddings
        self._setup_embeddings()
    
    def _setup_embeddings(self):
        """Configura embeddings usando HuggingFace"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception as e:
            st.error(f"Erro ao configurar embeddings: {str(e)}")
    
    def process_math_documents(self) -> bool:
        """Processa todos os documentos da pasta matemática"""
        if not os.path.exists(self.math_folder_path):
            st.error(f"Pasta não encontrada: {self.math_folder_path}")
            return False
        
        if not DOCUMENT_PROCESSING_AVAILABLE:
            st.error("Bibliotecas de processamento não disponíveis. Instale: pip install python-docx pypdf")
            return False
        
        try:
            # Lista todos os arquivos na pasta matemática
            math_files = []
            for root, dirs, files in os.walk(self.math_folder_path):
                for file in files:
                    if file.lower().endswith(('.docx', '.pdf', '.txt')):
                        math_files.append(os.path.join(root, file))
            
            st.info(f"Encontrados {len(math_files)} arquivos para processar...")
            
            # Processa cada arquivo
            all_documents = []
            progress_bar = st.progress(0)
            
            for i, file_path in enumerate(math_files):
                try:
                    docs = self._process_single_file(file_path)
                    all_documents.extend(docs)
                    progress_bar.progress((i + 1) / len(math_files))
                except Exception as e:
                    st.warning(f"Erro ao processar {file_path}: {str(e)}")
                    continue
            
            self.documents = all_documents
            
            if not self.documents:
                st.error("Nenhum documento foi processado com sucesso")
                return False
            
            # Cria ou atualiza vectorstore
            self._create_vectorstore()
            
            st.success(f"✅ Processados {len(self.documents)} chunks de {len(math_files)} arquivos")
            return True
            
        except Exception as e:
            st.error(f"Erro no processamento: {str(e)}")
            return False
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Processa um único arquivo"""
        file_extension = Path(file_path).suffix.lower()
        file_name = Path(file_path).name
        
        # Extrai conteúdo baseado no tipo
        if file_extension == '.docx':
            content = self._extract_docx_content(file_path)
        elif file_extension == '.pdf':
            content = self._extract_pdf_content(file_path)
        elif file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            return []
        
        if not content.strip():
            return []
        
        # Cria documento LangChain
        doc = Document(
            page_content=content,
            metadata={
                "source": file_name,
                "file_path": file_path,
                "type": "matematica",
                "topic": self._extract_topic_from_filename(file_name)
            }
        )
        
        # Divide em chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        chunks = text_splitter.split_documents([doc])
        return chunks
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivos DOCX"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            return "\n\n".join(content_parts)
        except Exception as e:
            st.warning(f"Erro ao processar DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extrai conteúdo de arquivos PDF"""
        try:
            reader = PdfReader(file_path)
            content_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    content_parts.append(text.strip())
            
            return "\n\n".join(content_parts)
        except Exception as e:
            st.warning(f"Erro ao processar PDF {file_path}: {str(e)}")
            return ""
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extrai tópico do nome do arquivo"""
        # Remove extensão
        topic = Path(filename).stem
        
        # Limpa caracteres especiais
        topic = topic.replace("_", " ").replace("-", " ")
        
        # Capitaliza primeira letra
        return topic.title()
    
    def _create_vectorstore(self):
        """Cria ou atualiza o vectorstore"""
        try:
            if not self.embeddings:
                raise Exception("Embeddings não configurados")
            
            # Cria vectorstore ChromaDB
            self.vectorstore = Chroma.from_documents(
                documents=self.documents,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
            
            # Configura retriever
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
        except Exception as e:
            st.error(f"Erro ao criar vectorstore: {str(e)}")
            raise
    
    def load_existing_vectorstore(self) -> bool:
        """Carrega vectorstore existente se disponível"""
        try:
            if os.path.exists(self.persist_directory) and self.embeddings:
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                
                self.retriever = self.vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}
                )
                
                # Testa se a vectorstore tem conteúdo
                try:
                    test_docs = self.retriever.invoke("teste")
                    print(f"✅ VectorStore carregada com {len(test_docs)} documentos de teste")
                    
                    # Simula documents para estatísticas
                    # Como não temos acesso direto aos documentos originais,
                    # vamos buscar uma amostra para estatísticas
                    sample_docs = self.vectorstore.similarity_search("matemática", k=100)
                    self.documents = sample_docs
                    print(f"📊 Amostra carregada: {len(sample_docs)} chunks")
                    
                except Exception as e:
                    print(f"⚠️ Erro no teste da VectorStore: {e}")
                
                return True
        except Exception as e:
            if 'st' in globals():
                st.warning(f"Não foi possível carregar vectorstore existente: {str(e)}")
            else:
                print(f"Não foi possível carregar vectorstore existente: {str(e)}")
        
        return False
    
    def create_rag_chain(self, api_key: str):
        """Cria chain RAG conversacional"""
        if not self.retriever:
            raise Exception("Retriever não configurado")
        
        # Configura LLM
        llm = GroqLLM(api_key=api_key)
        
        # Configura memória
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )
        
        # Cria chain conversacional
        self.rag_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=self.retriever,
            memory=self.memory,
            return_source_documents=True,
            verbose=False
        )
        
        # Personaliza o prompt
        self.rag_chain.combine_docs_chain.llm_chain.prompt.template = """
Você é o Professor Carlos, especialista em matemática do ENEM e professor particular da Sther Souza.

CONTEXTO RELEVANTE:
{context}

PERGUNTA DA STHER: {question}

INSTRUÇÕES:
- Use o contexto fornecido dos materiais de matemática para responder
- Seja didático e encorajador
- Use linguagem adequada para uma estudante de 17 anos
- Sempre elogie o potencial da Sther
- Inclua exemplos práticos quando possível
- Se não souber algo específico, seja honesto
- Foque sempre no ENEM

FÓRMULAS MATEMÁTICAS - USE SEMPRE LaTeX:
- Fórmulas inline: $f(x) = ax^2 + bx + c$ 
- Fórmulas em destaque: $$\\frac{{-b \\pm \\sqrt{{b^2-4ac}}}}{{2a}}$$
- Matrizes: $$\\begin{{pmatrix}} a & b \\\\ c & d \\end{{pmatrix}}$$
- Determinantes: $$\\det(A) = \\sum_{{j=1}}^{{n}} a_{{ij}} \\cdot C_{{ij}}$$
- Frações: $$\\frac{{numerador}}{{denominador}}$$
- Raízes: $$\\sqrt{{expressão}}$$
- Expoentes: $x^2$, $a^{{n+1}}$
- Índices: $a_{{ij}}$, $x_1$

SEMPRE use LaTeX entre $ ou $$ para matemática - isso melhora MUITO a legibilidade!

RESPOSTA:
"""
    
    def get_response(self, question: str) -> Dict[str, Any]:
        """Gera resposta usando RAG"""
        if not self.rag_chain:
            return {
                "answer": "Sistema RAG não inicializado. Configure sua API key.",
                "source_documents": []
            }
        
        try:
            result = self.rag_chain({
                "question": question,
                "chat_history": []
            })
            
            return result
            
        except Exception as e:
            return {
                "answer": f"Erro ao gerar resposta: {str(e)}",
                "source_documents": []
            }
    
    def search_relevant_content(self, query: str, k: int = 3) -> List[Document]:
        """Busca conteúdo relevante"""
        if not self.retriever:
            return []
        
        try:
            docs = self.retriever.get_relevant_documents(query)
            return docs[:k]
        except Exception as e:
            st.error(f"Erro na busca: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas da base de conhecimento"""
        stats = {
            "total_documentos": len(self.documents),
            "vectorstore_inicializado": self.vectorstore is not None,
            "retriever_configurado": self.retriever is not None,
            "pasta_matematica": self.math_folder_path,
            "diretorio_persistencia": self.persist_directory
        }
        
        # Conta tópicos únicos
        topics = set()
        for doc in self.documents:
            topic = doc.metadata.get("topic", "Desconhecido")
            topics.add(topic)
        
        stats["topicos_unicos"] = len(topics)
        stats["lista_topicos"] = sorted(list(topics))
        
        return stats
    
    def clear_memory(self):
        """Limpa memória da conversa"""
        if self.memory:
            self.memory.clear()

# Instância global
local_math_rag = LocalMathRAG() 